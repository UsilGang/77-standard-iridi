#!/usr/bin/env python3
"""Deterministic Standard as Code CLI for the iRidi automation standard.

The CLI never writes to Google Docs. `capture-google` is read-only and creates an
immutable local baseline. All other commands operate on local artifacts.
"""

from __future__ import annotations

import argparse
import hashlib
import html
import json
import mimetypes
import os
import re
import shutil
import sys
import tempfile
import unicodedata
import urllib.request
import zipfile
import xml.etree.ElementTree as ET
from collections import defaultdict
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

try:
    import yaml
except ImportError as exc:  # pragma: no cover
    raise SystemExit("PyYAML is required. Install tooling/requirements.txt") from exc


TOOLING_DIR = Path(__file__).resolve().parent
WORKSPACE_DIR = TOOLING_DIR.parent
DEFAULT_SOURCE_DIR = WORKSPACE_DIR / "standard-src"
DEFAULT_BUILD_DIR = WORKSPACE_DIR / "build"
DEFAULT_RELEASES_DIR = WORKSPACE_DIR / "releases"
DEFAULT_MANIFEST_CANDIDATE = (
    WORKSPACE_DIR.parents[1]
    / "Artifacts"
    / "Machine_Readable"
    / "standard_book_manifest_candidate_v1.yaml"
)
SCHEMA_DIR = TOOLING_DIR / "schemas"
TABLE_LAYOUT_CONTRACT = TOOLING_DIR / "table_layout_contract.yaml"
SEMANTIC_ENRICHMENT_CONTRACT = TOOLING_DIR / "semantic_enrichment_contract.yaml"
TOKEN_RE = re.compile(r"[\w-]{2,}", re.UNICODE)
STOPWORDS = {
    "а", "без", "бы", "в", "во", "вот", "вы", "где", "да", "для", "до", "его", "ее", "если", "же", "за",
    "и", "из", "или", "их", "к", "как", "какие", "какой", "ко", "ли", "на", "над", "не", "но", "о", "об",
    "от", "по", "под", "при", "про", "с", "со", "так", "такой", "там", "то", "у", "уже", "что", "это",
    "the", "and", "for", "from", "how", "what", "where", "with",
}


def utc_now() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def atomic_write_bytes(path: Path, data: bytes) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(path.suffix + ".tmp")
    tmp.write_bytes(data)
    tmp.replace(path)


def atomic_write_text(path: Path, text: str) -> None:
    atomic_write_bytes(path, text.encode("utf-8"))


def normalize_zip_archive(path: Path) -> None:
    """Rewrite an OOXML archive with stable member order and timestamps."""
    temp = path.with_suffix(path.suffix + ".normalized")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temp, "w") as target:
        for original in sorted(source.infolist(), key=lambda item: item.filename):
            info = zipfile.ZipInfo(original.filename, date_time=(2000, 1, 1, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = original.external_attr
            info.create_system = original.create_system
            target.writestr(info, source.read(original.filename), compress_type=zipfile.ZIP_DEFLATED, compresslevel=9)
    temp.replace(path)


def read_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def write_json(path: Path, value: Any, *, compact: bool = False) -> None:
    if compact:
        text = json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    else:
        text = json.dumps(value, ensure_ascii=False, sort_keys=True, indent=2) + "\n"
    atomic_write_text(path, text)


def read_yaml(path: Path) -> Any:
    return yaml.safe_load(path.read_text(encoding="utf-8"))


def write_yaml(path: Path, value: Any) -> None:
    atomic_write_text(
        path,
        yaml.safe_dump(value, allow_unicode=True, sort_keys=False, width=120),
    )


def slugify(text: str, *, fallback: str = "item", limit: int = 80) -> str:
    value = re.sub(r"[^\w\s-]", "", text.lower(), flags=re.UNICODE)
    value = re.sub(r"[\s_]+", "-", value.strip()).strip("-")
    return (value or fallback)[:limit]


def stable_uid(prefix: str, *parts: str) -> str:
    raw = "|".join(parts)
    digest = hashlib.sha1(raw.encode("utf-8")).hexdigest()[:10]
    readable_source = slugify(parts[-1] if parts else prefix, fallback="unit", limit=44).replace("-", "_")
    readable = re.sub(r"[^a-z0-9_]+", "_", readable_source.lower()).strip("_") or "unit"
    return f"{prefix}_{readable}_{digest}"


RU_TRANSLIT = str.maketrans(
    {
        "а": "a", "б": "b", "в": "v", "г": "g", "д": "d", "е": "e", "ё": "e", "ж": "zh", "з": "z",
        "и": "i", "й": "y", "к": "k", "л": "l", "м": "m", "н": "n", "о": "o", "п": "p", "р": "r",
        "с": "s", "т": "t", "у": "u", "ф": "f", "х": "h", "ц": "ts", "ч": "ch", "ш": "sh", "щ": "sch",
        "ъ": "", "ы": "y", "ь": "", "э": "e", "ю": "yu", "я": "ya",
    }
)


def ascii_semantic_slug(text: str, *, fallback: str = "topic", limit: int = 64) -> str:
    value = unicodedata.normalize("NFKC", text).casefold().translate(RU_TRANSLIT)
    value = re.sub(r"https?://\S+", " ", value)
    value = re.sub(r"\[\[asset:([^]]+)\]\]", r"asset \1", value)
    value = re.sub(r"[^a-z0-9]+", "_", value).strip("_")
    return (value or fallback)[:limit].rstrip("_")


def display_number_from_title(title: str) -> str | None:
    match = re.match(r"^\s*(?:раздел\s+)?(\d+(?:\.\d+)*)\.?\s+", title, flags=re.I)
    return match.group(1) if match else None


def title_without_display_number(title: str) -> str:
    value = re.sub(r"^\s*(?:раздел\s+)?\d+(?:\.\d+)*\.?\s+", "", title, flags=re.I).strip()
    return value or title.strip()


def contains_private_use(text: str) -> bool:
    for char in text:
        code = ord(char)
        if 0xE000 <= code <= 0xF8FF or 0xF0000 <= code <= 0xFFFFD or 0x100000 <= code <= 0x10FFFD:
            return True
    return False


def logical_markdown_lines(content: str) -> list[str]:
    """Split Markdown only on real line endings.

    Google Docs uses vertical-tab/form-feed characters for soft line breaks
    inside table cells. ``str.splitlines`` treats those controls as physical
    rows and can silently split one Markdown table into several blocks.
    """
    return content.replace("\v", " ").replace("\f", " ").split("\n")


def markdown_without_initial_heading(content: str) -> str:
    lines = logical_markdown_lines(content)
    if lines and markdown_heading(lines[0].strip()):
        lines = lines[1:]
    return "\n".join(lines).strip()


def markdown_plain_text(content: str) -> str:
    value = re.sub(r"!\[[^]]*\]\([^)]+\)", " ", content)
    value = re.sub(r"\[([^]]+)\]\([^)]+\)", r"\1", value)
    value = re.sub(r"^\s{0,3}#{1,6}\s*", "", value, flags=re.M)
    value = re.sub(r"[|:*_~`>#]+", " ", value)
    return re.sub(r"\s+", " ", value).strip()


def parity_plain_text(content: str) -> str:
    """Normalize punctuation and markup for loss detection across render formats."""
    value = html.unescape(content).replace("\v", " ").replace("\f", " ").replace("_", " ")
    value = re.sub(r"[^\w%©°²]+", " ", value, flags=re.UNICODE)
    return re.sub(r"\s+", " ", value).strip().casefold()


def derived_summary(title: str, content: str, *, limit: int = 280) -> str:
    body = markdown_plain_text(markdown_without_initial_heading(content))
    if not body:
        return title
    sentence = re.split(r"(?<=[.!?])\s+", body, maxsplit=1)[0]
    return sentence[:limit].rstrip()


def topic_classification(title: str, content: str, next_title: str | None = None) -> str:
    if contains_private_use(title) or "\ufffd" in title:
        return "artifact"
    if re.fullmatch(r"\s*\[\[asset:[^]]+\]\]\s*", title) or re.fullmatch(r"https?://\S+", title.strip()):
        return "attachment"
    body = markdown_without_initial_heading(content)
    body_text = markdown_plain_text(body)
    body_images = bool(re.search(r"!\[[^]]*\]\([^)]+\)", body))
    if body_text or body_images:
        return "content"
    current_number = display_number_from_title(title)
    next_number = display_number_from_title(next_title or "")
    if current_number and next_number and next_number.startswith(current_number + "."):
        return "container"
    return "gap"


def topic_semantic_uid(domain: str, title: str, storage_uid: str, node_kind: str, used: set[str]) -> str:
    if node_kind == "attachment":
        object_match = re.search(r"asset:([^]]+)", title)
        base = "attachment_" + ascii_semantic_slug(object_match.group(1) if object_match else title, fallback="attachment")
    elif node_kind == "artifact":
        base = "artifact_" + hashlib.sha1(storage_uid.encode("utf-8")).hexdigest()[:8]
    elif storage_uid.startswith("std_topic_preamble_"):
        base = "overview"
    else:
        base = ascii_semantic_slug(title_without_display_number(title), fallback="topic")
    candidate = f"std_topic_{domain}_{base}"[:128].rstrip("_")
    if candidate in used:
        candidate = f"{candidate[:119].rstrip('_')}_{hashlib.sha1(storage_uid.encode('utf-8')).hexdigest()[:8]}"
    used.add(candidate)
    return candidate


def topic_alias_defaults(title: str) -> list[str]:
    clean = title_without_display_number(title).strip()
    if not clean or contains_private_use(clean) or clean.startswith("[[asset:") or re.fullmatch(r"https?://\S+", clean):
        return []
    aliases = [clean]
    for value in re.findall(r"\(([^)]+)\)", clean):
        if 2 <= len(value.strip()) <= 80:
            aliases.append(value.strip())
    return list(dict.fromkeys(aliases))


def topic_question_defaults(title: str, node_kind: str) -> list[str]:
    if node_kind not in {"content", "gap"}:
        return []
    clean = title_without_display_number(title).rstrip(".?").strip()
    if not clean:
        return []
    return [f"Что нужно знать про {clean}?", f"Как применяется {clean}?"]


def iter_tabs(tabs: Iterable[dict[str, Any]]) -> Iterable[dict[str, Any]]:
    for tab in tabs or []:
        yield tab
        yield from iter_tabs(tab.get("childTabs") or [])


def tab_properties(tab: dict[str, Any]) -> dict[str, Any]:
    return tab.get("tabProperties") or {}


def document_tab(tab: dict[str, Any]) -> dict[str, Any]:
    return tab.get("documentTab") or {}


def image_objects(doc: dict[str, Any]) -> Iterable[tuple[str, str, dict[str, Any], str]]:
    """Yield tab_id, object_id, embedded object and placement kind."""
    for tab in iter_tabs(doc.get("tabs") or []):
        props = tab_properties(tab)
        tab_id = str(props.get("tabId") or "unknown")
        dtab = document_tab(tab)
        for object_id, item in (dtab.get("inlineObjects") or {}).items():
            embedded = ((item or {}).get("inlineObjectProperties") or {}).get("embeddedObject") or {}
            if embedded.get("imageProperties"):
                yield tab_id, object_id, embedded, "inline"
        for object_id, item in (dtab.get("positionedObjects") or {}).items():
            embedded = ((item or {}).get("positionedObjectProperties") or {}).get("embeddedObject") or {}
            if embedded.get("imageProperties"):
                yield tab_id, object_id, embedded, "positioned"


def extension_for(content_type: str | None, source_uri: str | None) -> str:
    if content_type:
        ext = mimetypes.guess_extension(content_type.split(";", 1)[0].strip())
        if ext:
            return ".jpg" if ext == ".jpe" else ext
    if source_uri:
        suffix = Path(source_uri.split("?", 1)[0]).suffix.lower()
        if suffix in {".png", ".jpg", ".jpeg", ".gif", ".svg", ".webp"}:
            return suffix
    return ".bin"


def capture_google(args: argparse.Namespace) -> int:
    google_root = Path(args.google_mcp_root or os.environ.get("GOOGLE_MCP_ROOT", r"C:\AI\mcp-google"))
    src = google_root / "src"
    if src.is_dir():
        sys.path.insert(0, str(src))
    try:
        from google_mcp.auth import get_credentials  # type: ignore
        from googleapiclient.discovery import build  # type: ignore
    except ImportError as exc:
        raise SystemExit("Run capture-google with the mcp-google Python environment") from exc

    creds = get_credentials(allow_interactive=False)
    docs = build("docs", "v1", credentials=creds, cache_discovery=False)
    drive = build("drive", "v3", credentials=creds, cache_discovery=False)
    doc = docs.documents().get(documentId=args.document_id, includeTabsContent=True).execute()
    metadata = (
        drive.files()
        .get(
            fileId=args.document_id,
            supportsAllDrives=True,
            fields="id,name,mimeType,createdTime,modifiedTime,version,size,webViewLink,lastModifyingUser(displayName)",
        )
        .execute()
    )
    revision = str(doc.get("revisionId") or metadata.get("version") or "unknown")
    revision_key = hashlib.sha256(revision.encode("utf-8")).hexdigest()[:16]
    out = Path(args.output).resolve() / f"rev-{revision_key}"
    if out.exists() and not args.force:
        print(f"EXISTS {out}")
        return 0
    out.mkdir(parents=True, exist_ok=True)

    previous_assets: dict[str, dict[str, Any]] = {}
    previous_manifest_path = out / "baseline-manifest.json"
    if previous_manifest_path.exists():
        previous_manifest = read_json(previous_manifest_path)
        previous_assets = {
            str(row.get("object_id")): row
            for row in previous_manifest.get("assets", [])
            if row.get("status") == "downloaded" and row.get("path")
        }

    document_bytes = json.dumps(doc, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode("utf-8")
    atomic_write_bytes(out / "document.json", document_bytes)
    write_json(out / "drive-metadata.json", metadata)

    assets_dir = out / "assets"

    def capture_asset(item: tuple[str, str, dict[str, Any], str]) -> dict[str, Any]:
        tab_id, object_id, embedded, placement = item
        image = embedded.get("imageProperties") or {}
        content_uri = image.get("contentUri")
        source_uri = image.get("sourceUri")
        previous = previous_assets.get(object_id)
        if previous:
            previous_path = out / str(previous["path"])
            if previous_path.is_file() and sha256_file(previous_path) == previous.get("sha256"):
                return {**previous, "tab_id": tab_id, "placement": placement, "reused": True}
        row: dict[str, Any] = {
            "object_id": object_id,
            "tab_id": tab_id,
            "placement": placement,
            "title": embedded.get("title"),
            "description": embedded.get("description"),
            "source_uri": source_uri,
            "size": embedded.get("size"),
            "status": "missing_content_uri" if not content_uri else "pending",
        }
        if content_uri:
            try:
                req = urllib.request.Request(content_uri, headers={"User-Agent": "standard-book/1.0"})
                with urllib.request.urlopen(req, timeout=90) as response:
                    data = response.read()
                    content_type = response.headers.get("Content-Type")
                ext = extension_for(content_type, source_uri)
                target = assets_dir / f"{object_id}{ext}"
                atomic_write_bytes(target, data)
                row.update(
                    {
                        "status": "downloaded",
                        "path": target.relative_to(out).as_posix(),
                        "mime_type": content_type,
                        "bytes": len(data),
                        "sha256": sha256_bytes(data),
                    }
                )
            except Exception as exc:  # noqa: BLE001 - evidence must retain individual failures
                row.update({"status": "download_failed", "error": f"{type(exc).__name__}: {exc}"})
        return row

    image_items = list(image_objects(doc))
    with ThreadPoolExecutor(max_workers=max(1, args.workers)) as pool:
        # executor.map keeps the source order, making the manifest deterministic.
        asset_rows = list(pool.map(capture_asset, image_items))

    tabs = [
        {
            "tab_id": tab_properties(tab).get("tabId"),
            "title": tab_properties(tab).get("title"),
            "index": tab_properties(tab).get("index"),
        }
        for tab in iter_tabs(doc.get("tabs") or [])
    ]
    manifest = {
        "schema_version": "1.0",
        "baseline_uid": f"baseline_{args.document_id[:8]}_{revision_key}",
        "captured_at": utc_now(),
        "document_id": args.document_id,
        "title": doc.get("title") or metadata.get("name"),
        "revision_id": revision,
        "revision_key": revision_key,
        "modified_time": metadata.get("modifiedTime"),
        "document_json_sha256": sha256_bytes(document_bytes),
        "tabs": tabs,
        "assets": asset_rows,
        "asset_counts": dict(
            sorted(
                {
                    status: sum(1 for row in asset_rows if row.get("status") == status)
                    for status in {row.get("status") for row in asset_rows}
                }.items()
            )
        ),
    }
    write_json(out / "baseline-manifest.json", manifest)
    print(json.dumps({"baseline": str(out), "tabs": len(tabs), "assets": manifest["asset_counts"]}, ensure_ascii=False))
    return 0


def paragraph_block(paragraph: dict[str, Any], element: dict[str, Any]) -> dict[str, Any]:
    style = paragraph.get("paragraphStyle") or {}
    runs: list[dict[str, Any]] = []
    text_parts: list[str] = []
    inline_ids: list[str] = []
    for item in paragraph.get("elements") or []:
        if item.get("textRun"):
            run = item["textRun"]
            content = run.get("content") or ""
            text_parts.append(content)
            runs.append({"text": content, "style": run.get("textStyle") or {}})
        elif item.get("inlineObjectElement"):
            object_id = item["inlineObjectElement"].get("inlineObjectId")
            if object_id:
                inline_ids.append(object_id)
                text_parts.append(f"[[asset:{object_id}]]")
    return {
        "type": "paragraph",
        "start_index": element.get("startIndex"),
        "end_index": element.get("endIndex"),
        "style": style.get("namedStyleType") or "NORMAL_TEXT",
        "heading_id": style.get("headingId"),
        "bullet": paragraph.get("bullet"),
        "text": "".join(text_parts).rstrip("\n"),
        "runs": runs,
        "inline_object_ids": inline_ids,
        "raw": element,
    }


def structural_text(content: list[dict[str, Any]]) -> str:
    parts: list[str] = []
    for element in content or []:
        if element.get("paragraph"):
            parts.append(paragraph_block(element["paragraph"], element)["text"])
        elif element.get("table"):
            for row in element["table"].get("tableRows") or []:
                for cell in row.get("tableCells") or []:
                    parts.append(structural_text(cell.get("content") or []))
    return " ".join(p for p in parts if p).strip()


def table_block(table: dict[str, Any], element: dict[str, Any]) -> dict[str, Any]:
    rows: list[list[str]] = []
    for row in table.get("tableRows") or []:
        rows.append([structural_text(cell.get("content") or []) for cell in row.get("tableCells") or []])
    return {
        "type": "table",
        "start_index": element.get("startIndex"),
        "end_index": element.get("endIndex"),
        "rows": rows,
        "raw": element,
    }


def body_blocks(body: dict[str, Any]) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    for element in body.get("content") or []:
        if element.get("paragraph"):
            out.append(paragraph_block(element["paragraph"], element))
        elif element.get("table"):
            out.append(table_block(element["table"], element))
        elif element.get("tableOfContents"):
            out.append({"type": "table_of_contents", "raw": element})
        elif element.get("sectionBreak"):
            out.append({"type": "section_break", "raw": element})
        else:
            out.append({"type": "unknown", "raw": element})
    return out


def markdown_for_block(block: dict[str, Any], asset_paths: dict[str, str]) -> str:
    if block["type"] == "paragraph":
        text = block.get("text") or ""
        for object_id in block.get("inline_object_ids") or []:
            ref = asset_paths.get(object_id, f"missing/{object_id}")
            text = text.replace(f"[[asset:{object_id}]]", f"![{object_id}]({ref})")
        style = block.get("style") or "NORMAL_TEXT"
        if style.startswith("HEADING_"):
            try:
                level = max(1, min(6, int(style.split("_", 1)[1])))
            except ValueError:
                level = 2
            return f"{'#' * level} {text.strip()}" if text.strip() else ""
        if block.get("bullet") and text.strip():
            return f"- {text.strip()}"
        return text.strip()
    if block["type"] == "table":
        rows = block.get("rows") or []
        if not rows:
            return ""
        width = max(len(row) for row in rows)
        normalized = [row + [""] * (width - len(row)) for row in rows]
        normalized = [
            [
                re.sub(
                    r"\[\[asset:([^]]+)\]\]",
                    lambda match: f"![{match.group(1)}]({asset_paths.get(match.group(1), f'missing/{match.group(1)}')})",
                    cell,
                )
                for cell in row
            ]
            for row in normalized
        ]
        lines = ["| " + " | ".join(cell.replace("|", "\\|") for cell in normalized[0]) + " |"]
        lines.append("| " + " | ".join("---" for _ in range(width)) + " |")
        lines.extend("| " + " | ".join(cell.replace("|", "\\|") for cell in row) + " |" for row in normalized[1:])
        return "\n".join(lines)
    return ""


def load_section_map(path: Path | None) -> dict[str, dict[str, Any]]:
    if not path or not path.is_file():
        return {}
    data = read_yaml(path) or {}
    return {str(row.get("source_tab_id")): row for row in ((data.get("book") or {}).get("sections") or [])}


def split_topics(section_uid: str, section_title: str, blocks: list[dict[str, Any]], asset_paths: dict[str, str]) -> list[dict[str, Any]]:
    topics: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None
    boundary_occurrences: dict[str, int] = defaultdict(int)

    def flush() -> None:
        nonlocal current
        if not current:
            return
        content = "\n\n".join(x for x in current.pop("markdown") if x).strip()
        if content:
            current["content"] = content + "\n"
            topics.append(current)
        current = None

    for index, block in enumerate(blocks, start=1):
        style = block.get("style") if block.get("type") == "paragraph" else None
        is_boundary = style in {"HEADING_1", "HEADING_2"} and (block.get("text") or "").strip()
        if is_boundary:
            flush()
            title = (block.get("text") or "").strip()
            anchor = str(block.get("heading_id") or title)
            boundary_occurrences[anchor] += 1
            occurrence = boundary_occurrences[anchor]
            seed = f"{anchor}|{occurrence}"
            current = {
                "uid": stable_uid("std_topic", section_uid, seed, f"{title}-{occurrence}"),
                "title": title,
                "heading_id": block.get("heading_id"),
                "markdown": [markdown_for_block(block, asset_paths)],
            }
        else:
            if current is None:
                current = {
                    "uid": stable_uid("std_topic", section_uid, "preamble"),
                    "title": section_title,
                    "heading_id": None,
                    "markdown": [],
                }
            current["markdown"].append(markdown_for_block(block, asset_paths))
    flush()
    return topics


def extract(args: argparse.Namespace) -> int:
    baseline = Path(args.baseline).resolve()
    source_dir = Path(args.source_dir).resolve()
    if (source_dir / "book.yaml").exists() and not args.force:
        raise SystemExit(f"Source tree already exists: {source_dir}. Use --force only for a deliberate baseline regeneration.")
    if args.force and source_dir.exists():
        source_root = source_dir.resolve()
        for name in ("assets", "sections", "staging", "entities", "taxonomy"):
            target = (source_root / name).resolve()
            if target.parent != source_root:
                raise SystemExit(f"Unsafe generated-path reset refused: {target}")
            if target.exists():
                shutil.rmtree(target)
        for name in ("book.yaml", "import-report.json"):
            target = source_root / name
            if target.exists():
                target.unlink()
    manifest = read_json(baseline / "baseline-manifest.json")
    doc = read_json(baseline / "document.json")
    section_map = load_section_map(Path(args.manifest_candidate).resolve() if args.manifest_candidate else None)

    source_dir.mkdir(parents=True, exist_ok=True)
    assets_dir = source_dir / "assets"
    assets_dir.mkdir(exist_ok=True)
    asset_paths: dict[str, str] = {}
    asset_manifest: list[dict[str, Any]] = []
    for row in manifest.get("assets") or []:
        if row.get("status") != "downloaded":
            asset_manifest.append(row)
            continue
        src = baseline / row["path"]
        target = assets_dir / src.name
        if not target.exists() or sha256_file(target) != row.get("sha256"):
            shutil.copy2(src, target)
        asset_paths[row["object_id"]] = f"../../assets/{target.name}"
        item = dict(row)
        item["source_path"] = row["path"]
        item["path"] = target.relative_to(source_dir).as_posix()
        asset_manifest.append(item)
    write_yaml(source_dir / "assets" / "manifest.yaml", {"baseline_uid": manifest["baseline_uid"], "assets": asset_manifest})

    section_refs: list[str] = []
    section_summaries: list[dict[str, Any]] = []
    staging_refs: list[str] = []
    for tab in iter_tabs(doc.get("tabs") or []):
        props = tab_properties(tab)
        tab_id = str(props.get("tabId") or "unknown")
        title = str(props.get("title") or tab_id)
        mapped = section_map.get(tab_id) or {}
        is_buffer = title.strip().lower() == "буфер" or mapped.get("archetype") == "staging_only"
        uid = str(mapped.get("uid") or stable_uid("std_ch", tab_id, title))
        target = (source_dir / "staging" / "buffer") if is_buffer else (source_dir / "sections" / uid)
        target.mkdir(parents=True, exist_ok=True)
        blocks = body_blocks((document_tab(tab).get("body") or {}))
        with (target / "blocks.jsonl").open("w", encoding="utf-8", newline="\n") as fh:
            for block in blocks:
                fh.write(json.dumps(block, ensure_ascii=False, sort_keys=True) + "\n")
        markdown = "\n\n".join(x for x in (markdown_for_block(block, asset_paths) for block in blocks) if x).strip() + "\n"
        atomic_write_text(target / "legacy.md", markdown)
        write_json(target / "source-tab.json", document_tab(tab))

        if is_buffer:
            write_yaml(
                target / "staging.yaml",
                {
                    "uid": "std_staging_buffer",
                    "type": "staging",
                    "status": "legacy_migrated",
                    "publish": False,
                    "auto_migrate": False,
                    "source_tab_id": tab_id,
                    "baseline_uid": manifest["baseline_uid"],
                },
            )
            staging_refs.append("std_staging_buffer")
            continue

        topic_asset_paths = {key: value.replace("../../assets/", "../../../../assets/") for key, value in asset_paths.items()}
        topics = split_topics(uid, title, blocks, topic_asset_paths)
        topic_refs: list[str] = []
        for topic in topics:
            topic_uid = topic["uid"]
            topic_dir = target / "topics" / topic_uid
            topic_dir.mkdir(parents=True, exist_ok=True)
            content_path = topic_dir / "content.md"
            atomic_write_text(content_path, topic["content"])
            topic_meta = {
                "uid": topic_uid,
                "type": "topic",
                "title": topic["title"],
                "parent_uid": uid,
                "status": "legacy_migrated",
                "revision": 1,
                "digest": sha256_file(content_path),
                "privacy": "internal",
                "source_refs": [f"google_doc:{args.document_id or manifest['document_id']}#tab={tab_id}", f"baseline:{manifest['baseline_uid']}"],
                "change_refs": [],
                "introduced_in": None,
                "last_changed_in": None,
                "allowed_outputs": ["legacy-fidelity", "internal_review"],
                "summary": topic["title"],
                "aliases": [],
                "audiences": ["integrator", "designer", "sales", "presales", "training"],
                "jobs": ["learn", "explain", "select", "audit"],
                "domains": [slugify(title).replace("-", "_")],
                "lifecycle": [],
                "entity_refs": [],
                "rule_refs": [],
                "coverage_status": "documented",
                "content_ref": content_path.relative_to(source_dir).as_posix(),
                "legacy_heading_id": topic.get("heading_id"),
            }
            write_yaml(topic_dir / "topic.yaml", topic_meta)
            write_yaml(topic_dir / "rules.yaml", {"topic_uid": topic_uid, "rules": []})
            topic_refs.append(topic_uid)
        section_meta = {
            "uid": uid,
            "type": "section",
            "title": mapped.get("title") or title,
            "display_order": int(mapped.get("display_order") or props.get("index", 0) + 1),
            "display_number": mapped.get("display_number"),
            "archetype": mapped.get("archetype") or "unclassified",
            "topic_refs": topic_refs,
            "source_tab_id": tab_id,
            "status": "legacy_migrated",
            "revision": 1,
            "digest": sha256_file(target / "legacy.md"),
            "privacy": "internal",
            "source_refs": [f"google_doc:{manifest['document_id']}#tab={tab_id}", f"baseline:{manifest['baseline_uid']}"],
            "change_refs": [],
            "introduced_in": None,
            "last_changed_in": None,
            "allowed_outputs": ["legacy-fidelity", "internal_review"],
        }
        write_yaml(target / "section.yaml", section_meta)
        section_refs.append(uid)
        section_summaries.append({"uid": uid, "title": section_meta["title"], "topics": len(topic_refs), "blocks": len(blocks)})

    book = {
        "uid": "std_iridi",
        "type": "book",
        "title": manifest.get("title") or "Стандарт автоматизации iRidi",
        "status": "legacy_migrated",
        "revision": 1,
        "digest": None,
        "privacy": "internal",
        "source_refs": [f"baseline:{manifest['baseline_uid']}"],
        "change_refs": [],
        "introduced_in": None,
        "last_changed_in": None,
        "allowed_outputs": ["legacy-fidelity", "internal_review"],
        "release": None,
        "baseline_uid": manifest["baseline_uid"],
        "source_revision_id": manifest["revision_id"],
        "section_refs": section_refs,
        "staging_refs": staging_refs,
    }
    write_yaml(source_dir / "book.yaml", book)
    write_json(source_dir / "import-report.json", {"created_at": utc_now(), "baseline_uid": manifest["baseline_uid"], "sections": section_summaries, "staging": staging_refs})
    print(json.dumps({"source_dir": str(source_dir), "sections": len(section_refs), "staging": len(staging_refs)}, ensure_ascii=False))
    return 0


def collect_units(source_dir: Path) -> tuple[dict[str, Any], list[dict[str, Any]], list[dict[str, Any]], list[dict[str, Any]]]:
    book = read_yaml(source_dir / "book.yaml")
    sections = [read_yaml(path) for path in sorted((source_dir / "sections").glob("*/section.yaml"))]
    topics = [read_yaml(path) for path in sorted((source_dir / "sections").glob("*/topics/*/topic.yaml"))]
    rules: list[dict[str, Any]] = []
    for path in sorted((source_dir / "sections").glob("*/topics/*/rules.yaml")):
        data = read_yaml(path) or {}
        rules.extend(data.get("rules") or [])
    return book, sections, topics, rules


def load_semantic_enrichment_contract(path: Path = SEMANTIC_ENRICHMENT_CONTRACT) -> dict[str, Any]:
    return read_yaml(path) if path.is_file() else {}


def remediate_migration_cmd(args: argparse.Namespace) -> int:
    source_dir = Path(args.source_dir).resolve()
    contract_path = Path(args.contract).resolve()
    contract = load_semantic_enrichment_contract(contract_path)
    change_uid = str(contract.get("change_uid") or "std_change_migration_semantic_addressing")
    domain_contract = contract.get("domains") or {}
    overrides = contract.get("topic_overrides") or {}
    book = read_yaml(source_dir / "book.yaml")
    section_paths = {path.parent.name: path for path in (source_dir / "sections").glob("*/section.yaml")}
    used_semantic_uids: set[str] = set()
    changed_topic_uids: list[str] = []
    changed_section_uids: list[str] = []
    mappings: list[dict[str, Any]] = []
    node_counts: dict[str, int] = defaultdict(int)

    for section_uid in book.get("section_refs") or []:
        section_path = section_paths[section_uid]
        section = read_yaml(section_path)
        section_dir = section_path.parent
        domain_row = domain_contract.get(section_uid) or {}
        domain = str(domain_row.get("id") or section_uid.removeprefix("std_ch_"))
        topic_records: list[tuple[Path, dict[str, Any], str]] = []
        for storage_uid in section.get("topic_refs") or []:
            topic_path = section_dir / "topics" / storage_uid / "topic.yaml"
            topic = read_yaml(topic_path)
            content = (source_dir / topic["content_ref"]).read_text(encoding="utf-8")
            topic_records.append((topic_path, topic, content))

        prepared: list[dict[str, Any]] = []
        for index, (topic_path, topic, content) in enumerate(topic_records):
            next_title = str(topic_records[index + 1][1].get("title") or "") if index + 1 < len(topic_records) else None
            title = str(topic.get("title") or "")
            node_kind = topic_classification(title, content, next_title)
            override = overrides.get(str(topic["uid"])) or {}
            semantic_uid = str(override.get("semantic_uid") or topic.get("semantic_uid") or "")
            if semantic_uid:
                if semantic_uid in used_semantic_uids:
                    raise ValueError(f"Duplicate semantic UID in remediation input: {semantic_uid}")
                used_semantic_uids.add(semantic_uid)
            else:
                semantic_uid = topic_semantic_uid(domain, title, str(topic["uid"]), node_kind, used_semantic_uids)
            prepared.append(
                {
                    "path": topic_path,
                    "topic": topic,
                    "content": content,
                    "node_kind": node_kind,
                    "semantic_uid": semantic_uid,
                    "override": override,
                }
            )

        number_to_uid: dict[str, str] = {}
        last_content_uid: str | None = None
        for row in prepared:
            topic = row["topic"]
            title = str(topic.get("title") or "")
            node_kind = row["node_kind"]
            semantic_uid = row["semantic_uid"]
            display_number = display_number_from_title(title)
            semantic_parent_uid = section_uid
            if display_number:
                pieces = display_number.split(".")
                for length in range(len(pieces) - 1, 0, -1):
                    candidate = ".".join(pieces[:length])
                    if candidate in number_to_uid:
                        semantic_parent_uid = number_to_uid[candidate]
                        break
            attached_to_uid = None
            if node_kind == "attachment":
                attached_to_uid = last_content_uid
                semantic_parent_uid = attached_to_uid or section_uid
            elif node_kind != "artifact":
                if display_number:
                    number_to_uid[display_number] = semantic_uid
                if node_kind in {"content", "gap"}:
                    last_content_uid = semantic_uid

            override = row["override"]
            aliases = list(dict.fromkeys([*topic_alias_defaults(title), *(override.get("aliases") or [])]))
            questions = list(dict.fromkeys([*topic_question_defaults(title, node_kind), *(override.get("answers_questions") or [])]))
            coverage_status = {
                "content": "documented",
                "container": "out_of_scope",
                "gap": "gap",
                "attachment": "documented",
                "artifact": "out_of_scope",
            }[node_kind]
            updated = dict(topic)
            updated.update(
                {
                    "revision": max(2, int(topic.get("revision") or 1)),
                    "change_refs": list(dict.fromkeys([*(topic.get("change_refs") or []), change_uid])),
                    "semantic_uid": semantic_uid,
                    "legacy_uids": list(dict.fromkeys([str(topic["uid"]), *(topic.get("legacy_uids") or [])])),
                    "node_kind": node_kind,
                    "display_number": display_number,
                    "semantic_parent_uid": semantic_parent_uid,
                    "attached_to_uid": attached_to_uid,
                    "summary": derived_summary(title, row["content"]),
                    "aliases": aliases,
                    "answers_questions": questions,
                    "domains": [domain],
                    "consumer_applications": [
                        "reference_agent",
                        "sales_technical_qa",
                        "presales_auditor",
                        "training_agent",
                        "editorial_agent",
                        "gap_router",
                    ],
                    "coverage_status": coverage_status,
                    "publish": node_kind != "artifact",
                    "queryable": node_kind in {"content", "gap"},
                    "search_metadata_origin": "machine_derived_baseline_remediation",
                    "metadata_review_status": "needs_editorial_review",
                }
            )
            if updated != topic:
                write_yaml(row["path"], updated)
                changed_topic_uids.append(str(topic["uid"]))
            node_counts[node_kind] += 1
            mappings.append(
                {
                    "legacy_uid": topic["uid"],
                    "semantic_uid": semantic_uid,
                    "node_kind": node_kind,
                    "parent_uid": semantic_parent_uid,
                    "attached_to_uid": attached_to_uid,
                    "content_ref": topic["content_ref"],
                }
            )

        updated_section = dict(section)
        section_coverage = "documented" if prepared else "gap"
        updated_section.update(
            {
                "revision": max(2, int(section.get("revision") or 1)),
                "change_refs": list(dict.fromkeys([*(section.get("change_refs") or []), change_uid])),
                "semantic_uid": section_uid,
                "legacy_uids": list(dict.fromkeys([section_uid, *(section.get("legacy_uids") or [])])),
                "node_kind": "section",
                "domain": domain,
                "aliases": list(domain_row.get("aliases") or []),
                "coverage_status": section_coverage,
            }
        )
        if updated_section != section:
            write_yaml(section_path, updated_section)
            changed_section_uids.append(section_uid)

    updated_book = dict(book)
    updated_book.update(
        {
            "revision": max(2, int(book.get("revision") or 1)),
            "change_refs": list(dict.fromkeys([*(book.get("change_refs") or []), change_uid])),
            "semantic_uid": str(book.get("semantic_uid") or book.get("uid") or "std_iridi"),
        }
    )
    if updated_book != book:
        write_yaml(source_dir / "book.yaml", updated_book)

    assets = (read_yaml(source_dir / "assets" / "manifest.yaml") or {}).get("assets") or []
    unplaced_assets = [
        {
            "object_id": row.get("object_id"),
            "tab_id": row.get("tab_id"),
            "path": row.get("path"),
            "sha256": row.get("sha256"),
            "placement": "positioned",
            "placement_status": "needs_manual_placement",
        }
        for row in assets
        if row.get("placement") == "positioned"
    ]
    report = {
        "schema_version": "1.0",
        "status": "applied_to_local_migrated_source",
        "applied_at": utc_now(),
        "source_ref": contract.get("source_ref"),
        "change_uid": change_uid,
        "contract_ref": contract_path.relative_to(WORKSPACE_DIR).as_posix() if contract_path.is_relative_to(WORKSPACE_DIR) else str(contract_path),
        "counts": {
            "sections": len(book.get("section_refs") or []),
            "topics": len(mappings),
            "changed_sections": len(changed_section_uids),
            "changed_topics": len(changed_topic_uids),
            "node_kinds": dict(sorted(node_counts.items())),
            "unplaced_positioned_assets": len(unplaced_assets),
        },
        "mappings": mappings,
        "unplaced_assets": unplaced_assets,
        "semantic_content_changed": False,
    }
    output = Path(args.output).resolve() if args.output else source_dir / "transformation-manifest.yaml"
    write_yaml(output, report)
    print(json.dumps({key: report[key] for key in ("status", "change_uid", "counts", "semantic_content_changed")}, ensure_ascii=False))
    return 0


def validate(args: argparse.Namespace) -> int:
    source_dir = Path(args.source_dir).resolve()
    findings: list[dict[str, Any]] = []
    try:
        book, sections, topics, rules = collect_units(source_dir)
    except Exception as exc:  # noqa: BLE001
        findings.append({"severity": "error", "code": "load_failed", "message": str(exc)})
        book, sections, topics, rules = {}, [], [], []

    all_units = [book, *sections, *topics, *rules]
    seen: dict[str, str] = {}
    for unit in all_units:
        uid = str((unit or {}).get("uid") or "")
        if not uid:
            findings.append({"severity": "error", "code": "missing_uid"})
            continue
        if uid in seen:
            findings.append({"severity": "error", "code": "duplicate_uid", "uid": uid})
        seen[uid] = str((unit or {}).get("type") or "unknown")
    section_uids = {row.get("uid") for row in sections}
    for ref in book.get("section_refs") or []:
        if ref not in section_uids:
            findings.append({"severity": "error", "code": "missing_section_ref", "ref": ref})
    topic_uids = {row.get("uid") for row in topics}
    for section in sections:
        refs = section.get("topic_refs") or []
        if len(refs) != len(set(refs)):
            findings.append({"severity": "error", "code": "duplicate_topic_ref", "section": section.get("uid")})
        for ref in section.get("topic_refs") or []:
            if ref not in topic_uids:
                findings.append({"severity": "error", "code": "missing_topic_ref", "section": section.get("uid"), "ref": ref})
    referenced_topic_uids = {ref for section in sections for ref in (section.get("topic_refs") or [])}
    for orphan in sorted(topic_uids - referenced_topic_uids):
        findings.append({"severity": "error", "code": "orphan_topic", "uid": orphan})
    for topic in topics:
        content = source_dir / str(topic.get("content_ref") or "")
        if not content.is_file():
            findings.append({"severity": "error", "code": "missing_content", "uid": topic.get("uid"), "path": str(content)})
        elif topic.get("digest") and topic["digest"] != sha256_file(content):
            findings.append({"severity": "error", "code": "digest_mismatch", "uid": topic.get("uid")})
    asset_manifest = source_dir / "assets" / "manifest.yaml"
    if asset_manifest.is_file():
        for row in (read_yaml(asset_manifest) or {}).get("assets") or []:
            if row.get("status") == "downloaded":
                path = source_dir / row.get("path", "")
                if not path.is_file():
                    findings.append({"severity": "error", "code": "missing_asset", "object_id": row.get("object_id")})
                elif row.get("sha256") and sha256_file(path) != row["sha256"]:
                    findings.append({"severity": "error", "code": "asset_digest_mismatch", "object_id": row.get("object_id")})

    schema_errors: list[str] = []
    try:
        import jsonschema

        schema = read_json(SCHEMA_DIR / "standard.schema.json")
        for unit in all_units:
            try:
                jsonschema.validate(unit, schema)
            except jsonschema.ValidationError as exc:
                schema_errors.append(f"{unit.get('uid', '<missing>')}: {exc.message}")
    except ImportError:
        findings.append({"severity": "warning", "code": "jsonschema_not_installed"})
    findings.extend({"severity": "error", "code": "schema", "message": msg} for msg in schema_errors)
    report = {
        "status": "pass" if not any(x["severity"] == "error" for x in findings) else "fail",
        "checked_at": utc_now(),
        "counts": {"sections": len(sections), "topics": len(topics), "rules": len(rules)},
        "findings": findings,
    }
    out = Path(args.output).resolve() if args.output else DEFAULT_BUILD_DIR / "validation-report.json"
    write_json(out, report)
    print(json.dumps(report, ensure_ascii=False))
    return 0 if report["status"] == "pass" else 2


def ordered_section_topics(section: dict[str, Any], topics_by_uid: dict[str, dict[str, Any]]) -> list[dict[str, Any]]:
    return [topics_by_uid[uid] for uid in (section.get("topic_refs") or []) if uid in topics_by_uid]


def selected_section_refs(book: dict[str, Any], sections: list[dict[str, Any]], requested: list[str] | None = None) -> list[str]:
    """Return requested section UIDs in book order, or the complete book scope."""
    book_refs = list(book.get("section_refs") or [])
    if not requested:
        return book_refs
    known = {str(section.get("uid")) for section in sections}
    unknown = sorted(set(requested) - known)
    outside_book = sorted(set(requested) - set(book_refs))
    if unknown:
        raise SystemExit(f"Unknown section UID(s): {', '.join(unknown)}")
    if outside_book:
        raise SystemExit(f"Section UID(s) are not in book.section_refs: {', '.join(outside_book)}")
    requested_set = set(requested)
    return [ref for ref in book_refs if ref in requested_set]


def markdown_table_cells(line: str) -> list[str]:
    value = line.strip().strip("|")
    return [cell.strip().replace("\\|", "|") for cell in re.split(r"(?<!\\)\|", value)]


def markdown_is_separator(cells: list[str]) -> bool:
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def load_table_layout_contract(path: Path = TABLE_LAYOUT_CONTRACT) -> dict[str, dict[int, dict[str, Any]]]:
    if not path.is_file():
        return {}
    data = read_yaml(path) or {}
    result: dict[str, dict[int, dict[str, Any]]] = defaultdict(dict)
    for row in data.get("overrides") or []:
        content_ref = str(row.get("content_ref") or "")
        table_index = int(row.get("table_index", -1))
        if content_ref and table_index >= 0:
            result[content_ref][table_index] = row
    return dict(result)


def table_cell_plain_text(cell: str) -> str:
    value = re.sub(r"!\[[^]]*\]\([^)]+\)", "", cell)
    return re.sub(r"https?://\S+", "", value).strip()


def infer_table_header_rows(rows: list[list[str]]) -> int:
    if len(rows) <= 1:
        return 0
    first = rows[0]
    if any("![" in cell for cell in first):
        return 0
    if any(len(table_cell_plain_text(cell)) > 80 for cell in first):
        return 0
    if any(not table_cell_plain_text(cell) for cell in first):
        return 0
    header_terms = {
        "базовое сечение",
        "возможная интеграция",
        "где применять",
        "код продукта",
        "контролируемый параметр",
        "критерий",
        "линейка lite",
        "линейка pro",
        "максимальный пусковой (ударный) ток",
        "максимальный ток",
        "марка кабеля",
        "модель",
        "назначение",
        "номинальный ток",
        "общий свет",
        "определение",
        "особенности",
        "параметр",
        "подустройство",
        "поддерживаемые ассистенты",
        "преимущества",
        "привод / блок управления",
        "применение",
        "производитель",
        "пусковой (ударный) ток",
        "рекомендуемое значение",
        "рекомендуемый кабель",
        "светильник",
        "серия берлин",
        "сцена",
        "сценарий",
        "схема",
        "температура света",
        "термин",
        "тип датчика",
        "тип светильника",
        "тип устройства",
        "функционал",
        "функция",
        "шторы",
        "№",
    }
    normalized_first = {re.sub(r"\s+", " ", table_cell_plain_text(cell)).strip().casefold() for cell in first}
    return 1 if normalized_first & header_terms else 0


def infer_column_width_percent(rows: list[list[str]]) -> list[float]:
    width = max((len(row) for row in rows), default=1)
    if width <= 1:
        return [100.0]
    image_columns = [any("![" in row[index] for row in rows) for index in range(width)]
    if width == 2 and image_columns == [False, True]:
        return [60.0, 40.0]
    if width == 2 and image_columns == [True, False]:
        return [40.0, 60.0]
    scores: list[float] = []
    for index in range(width):
        lengths = [len(table_cell_plain_text(row[index])) for row in rows]
        text_score = max(lengths, default=0) ** 0.5 * 5
        scores.append(max(18.0, min(60.0, text_score) + (28.0 if image_columns[index] else 0.0)))
    total = sum(scores) or 1.0
    return [round(score * 100.0 / total, 3) for score in scores]


def analyze_table_rows(rows: list[list[str]], override: dict[str, Any] | None = None) -> tuple[list[list[str]], int, list[float]]:
    override = override or {}
    width = max((len(row) for row in rows), default=1)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    explicit_width = int(override.get("effective_columns") or 0)
    if explicit_width:
        width = max(1, min(explicit_width, width))
    else:
        while width > 1 and all(not row[width - 1].strip() for row in normalized):
            width -= 1
    normalized = [row[:width] for row in normalized]
    header_rows = int(override["header_rows"]) if "header_rows" in override else infer_table_header_rows(normalized)
    widths = [float(value) for value in (override.get("column_width_percent") or infer_column_width_percent(normalized))]
    if len(widths) != width or any(value <= 0 for value in widths):
        raise ValueError(f"Invalid table column width contract: expected {width} positive values, got {widths}")
    total = sum(widths)
    widths = [round(value * 100.0 / total, 3) for value in widths]
    return normalized, max(0, min(header_rows, len(normalized))), widths


def legacy_table_rows(rows: list[list[str]]) -> tuple[list[list[str]], int, list[float]]:
    width = max((len(row) for row in rows), default=1)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    return normalized, min(1, len(normalized)), [round(100.0 / width, 3)] * width


def set_docx_table_geometry(table: Any, column_width_percent: list[float], total_width_dxa: int = 9072) -> list[int]:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    widths = [round(total_width_dxa * value / 100.0) for value in column_width_percent]
    widths[-1] += total_width_dxa - sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for tag, attrs in (
        ("w:tblW", {"w:type": "dxa", "w:w": str(total_width_dxa)}),
        ("w:tblLayout", {"w:type": "fixed"}),
        ("w:tblInd", {"w:type": "dxa", "w:w": "0"}),
    ):
        element = tbl_pr.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tbl_pr.append(element)
        for key, value in attrs.items():
            element.set(qn(key), value)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[index]))
    return widths


def markdown_heading(line: str) -> tuple[int, str] | None:
    match = re.fullmatch(r"(#{1,6})\s+(.+?)\s*", line.strip())
    if not match:
        return None
    return len(match.group(1)), match.group(2)


def rendered_heading_level(markdown_level: int, normalized: bool) -> int:
    """Normalize all headings inside an addressable topic to one subheading level."""
    if normalized:
        return 3
    return 2 if markdown_level <= 2 else 3


HTML_CSS = (
    "html{background:#F1F4F6}"
    "body{font-family:Arial,sans-serif;margin:0;font-size:16px;line-height:1.55;color:#202124;overflow-wrap:anywhere}"
    ".book-shell{display:grid;grid-template-columns:minmax(260px,320px) minmax(0,1040px);gap:28px;"
    "max-width:1420px;margin:0 auto;padding:24px;box-sizing:border-box;align-items:start}"
    ".book-content{min-width:0;padding:48px clamp(28px,6vw,80px) 96px;box-sizing:border-box;background:#FFF;"
    "box-shadow:0 2px 14px rgba(21,58,91,.08)}"
    ".toc-panel{position:sticky;top:24px;max-height:calc(100vh - 48px);overflow:auto;background:#FFF;"
    "border:1px solid #D5DEE5;border-radius:10px;padding:20px;box-sizing:border-box}"
    ".toc-panel h2{margin:0 0 .75rem;font-size:1.25rem}.toc-panel ol{list-style:none;margin:.35rem 0;padding-left:1.2rem}"
    ".toc-panel li{margin:.35rem 0}.toc-panel details>ol{margin-top:.55rem}.toc-panel summary{cursor:pointer;font-weight:700;color:#153A5B}"
    ".toc-panel a{color:#244E70;text-decoration:none}.toc-panel a:hover,.toc-panel a:focus{text-decoration:underline}"
    ".toc-section-start{display:block;margin:.45rem 0 .65rem;font-size:.9rem}.toc-fragments{font-size:.88rem;color:#4B5563}"
    ".topic-navigation{display:flex;gap:.75rem;justify-content:space-between;align-items:center;flex-wrap:wrap;"
    "margin:2rem 0 .5rem;padding:.55rem .75rem;border-top:1px solid #D5DEE5;border-bottom:1px solid #D5DEE5;font-size:.9rem}"
    ".topic-navigation a,.back-to-toc{color:#244E70;text-decoration:none}.topic-navigation a:hover,.back-to-toc:hover{text-decoration:underline}"
    ".back-to-toc{position:fixed;right:24px;bottom:24px;background:#153A5B;color:#FFF;padding:.65rem .9rem;border-radius:999px;"
    "box-shadow:0 2px 10px rgba(0,0,0,.2);z-index:5}.back-to-toc:hover{color:#FFF}"
    "h1,h2,h3{color:#153A5B;line-height:1.25}"
    "h1{margin:0 0 1.5rem}h2{margin:2.25rem 0 .9rem}h3{margin:1.75rem 0 .7rem}"
    "p{margin:.7rem 0 1rem}ul{margin:.6rem 0 1.2rem;padding-left:1.75rem}li{margin:.25rem 0}"
    "img{display:block;max-width:100%;height:auto;margin:1rem auto}figure{margin:1.25rem 0}"
    "table{width:100%;table-layout:fixed;border-collapse:collapse;margin:1.25rem 0;font-size:.95rem}"
    "td,th{border:1px solid #B8C4CE;padding:12px;vertical-align:top;overflow-wrap:anywhere}th{background:#EAF1F6}"
    "td.table-cell-image{font-size:.85rem;color:#4B5563}td.table-cell-image img{width:auto;max-width:100%;max-height:320px;object-fit:contain}"
    ".missing-asset{color:#9B1C1C}.gap-note,.unplaced-assets{border-left:4px solid #C98200;background:#FFF8E6;padding:10px 14px;margin:1rem 0}"
    ".machine-anchor{scroll-margin-top:24px}"
    "@media(max-width:1000px){.book-shell{display:block;padding:0}.toc-panel{position:relative;top:auto;max-height:none;margin:16px;}.book-content{box-shadow:none}}"
    "@media(max-width:640px){body{font-size:15px}.book-content{padding:28px 22px 64px}h1{font-size:1.75rem}h2{font-size:1.4rem}.back-to-toc{right:14px;bottom:14px}}"
)


def fragment_anchor_uid(topic_uid: str, title: str, occurrence: int = 1) -> str:
    """Return a deterministic, HTML-safe anchor for a heading inside a topic."""
    prefix = f"std_fragment_{topic_uid.removeprefix('std_topic_')}"
    base = f"{prefix}_{hashlib.sha1(title.encode('utf-8')).hexdigest()[:10]}"
    return base if occurrence == 1 else f"{base}_{occurrence}"


def markdown_navigation_headings(path: Path, topic_uid: str, skip_initial_heading: str | None = None) -> list[dict[str, str]]:
    """Extract addressable subheadings using the same rules as the renderer."""
    result: list[dict[str, str]] = []
    checked_initial_heading = False
    occurrences: dict[str, int] = defaultdict(int)
    for line in logical_markdown_lines(path.read_text(encoding="utf-8")):
        stripped = line.strip()
        if not stripped:
            continue
        heading = markdown_heading(stripped)
        if not checked_initial_heading:
            checked_initial_heading = True
            if skip_initial_heading and heading and heading[1] == skip_initial_heading:
                continue
        if not heading or re.search(r"!\[([^]]*)\]\(([^)]+)\)", heading[1]):
            continue
        occurrences[heading[1]] += 1
        result.append(
            {
                "uid": fragment_anchor_uid(topic_uid, heading[1], occurrences[heading[1]]),
                "title": heading[1],
            }
        )
    return result


def markdown_cell_html(cell: str, source_path: Path) -> str:
    parts: list[str] = []
    cursor = 0
    for match in re.finditer(r"!\[([^]]*)\]\(([^)]+)\)", cell):
        parts.append(html.escape(cell[cursor:match.start()]))
        image_path = (source_path.parent / match.group(2)).resolve()
        if image_path.is_file():
            parts.append(f"<img src='assets/{html.escape(image_path.name)}' alt='{html.escape(match.group(1))}'>")
        else:
            parts.append(html.escape(match.group(0)))
        cursor = match.end()
    parts.append(html.escape(cell[cursor:]))
    return "".join(parts)


def paragraph_style_hint(block: dict[str, Any]) -> dict[str, Any] | None:
    if block.get("type") != "paragraph":
        return None
    segments: list[dict[str, Any]] = []
    for run in block.get("runs") or []:
        text = str(run.get("text") or "").rstrip("\n").replace("\v", " ").replace("\f", " ")
        if not text:
            continue
        style = run.get("style") or {}
        segments.append(
            {
                "text": text,
                "bold": bool(style.get("bold")),
                "italic": bool(style.get("italic")),
                "underline": bool(style.get("underline")),
                "link": str((style.get("link") or {}).get("url") or ""),
            }
        )
    if not segments:
        return None
    text = "".join(str(segment.get("text") or "") for segment in segments).strip()
    return {"text": text, "segments": segments}


def section_topic_style_hints(source_dir: Path, section: dict[str, Any]) -> dict[str, dict[str, list[dict[str, Any]]]]:
    blocks_path = source_dir / "sections" / str(section["uid"]) / "blocks.jsonl"
    if not blocks_path.is_file():
        return {}
    blocks = [json.loads(line) for line in blocks_path.read_text(encoding="utf-8").split("\n") if line.strip()]
    groups: list[list[dict[str, Any]]] = []
    current: list[dict[str, Any]] = []
    for block in blocks:
        boundary = (
            block.get("type") == "paragraph"
            and block.get("style") in {"HEADING_1", "HEADING_2"}
            and str(block.get("text") or "").strip()
        )
        if boundary and current:
            groups.append(current)
            current = []
        current.append(block)
    if current:
        groups.append(current)
    groups = [group for group in groups if any(block.get("type") in {"paragraph", "table"} for block in group)]
    refs = list(section.get("topic_refs") or [])
    if len(groups) != len(refs):
        return {}
    result: dict[str, dict[str, list[dict[str, Any]]]] = {}
    for topic_uid, group in zip(refs, groups):
        by_text: dict[str, list[dict[str, Any]]] = defaultdict(list)
        for block in group:
            hint = paragraph_style_hint(block)
            if hint:
                by_text[parity_plain_text(hint["text"])].append(hint)
        result[str(topic_uid)] = dict(by_text)
    return result


def take_style_hint(style_hints: dict[str, list[dict[str, Any]]] | None, text: str) -> dict[str, Any] | None:
    if not style_hints:
        return None
    key = parity_plain_text(markdown_plain_text(text))
    candidates = style_hints.get(key) or []
    return candidates.pop(0) if candidates else None


def styled_html(text: str, hint: dict[str, Any] | None) -> str:
    if not hint or parity_plain_text(hint.get("text") or "") != parity_plain_text(markdown_plain_text(text)):
        return html.escape(text)
    parts: list[str] = []
    for segment in hint.get("segments") or []:
        value = html.escape(str(segment.get("text") or ""))
        if segment.get("bold"):
            value = f"<strong>{value}</strong>"
        if segment.get("italic"):
            value = f"<em>{value}</em>"
        if segment.get("underline"):
            value = f"<u>{value}</u>"
        link = str(segment.get("link") or "")
        if link.startswith(("http://", "https://", "mailto:")):
            value = f"<a href='{html.escape(link, quote=True)}'>{value}</a>"
        parts.append(value)
    return "".join(parts) or html.escape(text)


def styled_image_paragraph_html(text: str, source_path: Path, hint: dict[str, Any] | None) -> str:
    parts: list[str] = []
    cursor = 0
    image_pattern = re.compile(r"!\[([^]]*)\]\(([^)]+)\)")
    for match in image_pattern.finditer(text):
        plain = text[cursor:match.start()].strip()
        if plain:
            parts.append(styled_html(plain, hint if parity_plain_text(plain) == parity_plain_text((hint or {}).get("text") or "") else None))
        image_path = (source_path.parent / match.group(2)).resolve()
        if image_path.is_file():
            parts.append(f"<img src='assets/{html.escape(image_path.name)}' alt='{html.escape(match.group(1))}'>")
        else:
            parts.append(html.escape(match.group(0)))
        cursor = match.end()
    plain = text[cursor:].strip()
    if plain:
        parts.append(styled_html(plain, hint if parity_plain_text(plain) == parity_plain_text((hint or {}).get("text") or "") else None))
    return "".join(parts)


def markdown_html(
    path: Path,
    skip_initial_heading: str | None = None,
    table_layouts: dict[int, dict[str, Any]] | None = None,
    suppress_initial_heading: bool = False,
    fragment_uid_prefix: str | None = None,
    style_hints: dict[str, list[dict[str, Any]]] | None = None,
) -> list[str]:
    lines = logical_markdown_lines(path.read_text(encoding="utf-8"))
    out: list[str] = []
    index = 0
    table_index = 0
    checked_initial_heading = False
    heading_occurrences: dict[str, int] = defaultdict(int)
    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped:
            index += 1
            continue
        if not checked_initial_heading:
            checked_initial_heading = True
            heading = markdown_heading(stripped)
            if suppress_initial_heading and heading:
                stripped = heading[1]
            elif skip_initial_heading and heading and heading[1] == skip_initial_heading:
                take_style_hint(style_hints, heading[1])
                index += 1
                continue
        if stripped.startswith("| "):
            rows: list[list[str]] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                cells = markdown_table_cells(lines[index])
                if not markdown_is_separator(cells):
                    rows.append(cells)
                index += 1
            if rows:
                if skip_initial_heading is not None:
                    rows, header_rows, column_widths = analyze_table_rows(rows, (table_layouts or {}).get(table_index))
                else:
                    rows, header_rows, column_widths = legacy_table_rows(rows)
                out.append("<table class='table-headered'>" if header_rows else "<table class='table-headerless'>")
                out.append("<colgroup>" + "".join(f"<col style='width:{width:g}%'>" for width in column_widths) + "</colgroup>")
                if header_rows:
                    out.append("<thead>")
                    for row in rows[:header_rows]:
                        out.append("<tr>" + "".join(f"<th>{markdown_cell_html(cell, path)}</th>" for cell in row) + "</tr>")
                    out.append("</thead>")
                out.append("<tbody>")
                for row in rows[header_rows:]:
                    cells_html = []
                    for cell in row:
                        cell_class = " class='table-cell-image'" if "![" in cell else ""
                        cells_html.append(f"<td{cell_class}>{markdown_cell_html(cell, path)}</td>")
                    out.append("<tr>" + "".join(cells_html) + "</tr>")
                out.append("</tbody></table>")
                table_index += 1
            continue
        heading = markdown_heading(stripped)
        if heading and re.search(r"!\[([^]]*)\]\(([^)]+)\)", heading[1]):
            out.append(f"<div class='heading-asset'>{markdown_cell_html(heading[1], path)}</div>")
        elif re.search(r"!\[([^]]*)\]\(([^)]+)\)", stripped):
            out.append(f"<p>{styled_image_paragraph_html(stripped, path, take_style_hint(style_hints, stripped))}</p>")
        elif heading:
            level = rendered_heading_level(heading[0], normalized=skip_initial_heading is not None)
            fragment_attr = ""
            if fragment_uid_prefix:
                heading_occurrences[heading[1]] += 1
                topic_uid = fragment_uid_prefix.removeprefix("std_fragment_")
                fragment_uid = fragment_anchor_uid(f"std_topic_{topic_uid}", heading[1], heading_occurrences[heading[1]])
                fragment_attr = f" id='{html.escape(fragment_uid)}' data-fragment-uid='{html.escape(fragment_uid)}'"
            take_style_hint(style_hints, heading[1])
            out.append(f"<h{level}{fragment_attr}>{html.escape(heading[1])}</h{level}>")
        elif stripped.startswith("- "):
            items: list[str] = []
            while index < len(lines):
                candidate = lines[index].strip()
                if candidate.startswith("- "):
                    item_text = candidate[2:]
                    hint = take_style_hint(style_hints, item_text)
                    items.append(markdown_cell_html(item_text, path) if "![" in item_text else styled_html(item_text, hint))
                    index += 1
                    continue
                if not candidate:
                    probe = index + 1
                    while probe < len(lines) and not lines[probe].strip():
                        probe += 1
                    if probe < len(lines) and lines[probe].strip().startswith("- "):
                        index = probe
                        continue
                break
            out.append("<ul>" + "".join(f"<li>{item}</li>" for item in items) + "</ul>")
            continue
        else:
            out.append(f"<p>{styled_html(stripped, take_style_hint(style_hints, stripped))}</p>")
        index += 1
    return out


def topic_public_uid(topic: dict[str, Any]) -> str:
    return str(topic.get("semantic_uid") or topic.get("uid"))


def positioned_assets_by_tab(source_dir: Path) -> dict[str, list[dict[str, Any]]]:
    path = source_dir / "assets" / "manifest.yaml"
    result: dict[str, list[dict[str, Any]]] = defaultdict(list)
    if not path.is_file():
        return {}
    for row in (read_yaml(path) or {}).get("assets") or []:
        if row.get("placement") == "positioned" and row.get("status") == "downloaded":
            result[str(row.get("tab_id"))].append(row)
    return dict(result)


def navigation_outline(
    source_dir: Path,
    book: dict[str, Any],
    sections: list[dict[str, Any]],
    topics: list[dict[str, Any]],
    profile: str,
    section_uids: list[str] | None = None,
) -> list[dict[str, Any]]:
    """Build the shared navigation model used by HTML and DOCX."""
    section_by_uid = {row["uid"]: row for row in sections}
    topics_by_uid = {row["uid"]: row for row in topics}
    outline: list[dict[str, Any]] = []
    for ref in selected_section_refs(book, sections, section_uids):
        section = section_by_uid[ref]
        section_row: dict[str, Any] = {
            "uid": str(section.get("semantic_uid") or section["uid"]),
            "title": f"{section.get('display_number') or ''} {section['title']}".strip(),
            "topics": [],
        }
        if profile == "standard-normalized":
            for topic in ordered_section_topics(section, topics_by_uid):
                node_kind = str(topic.get("node_kind") or "content")
                if node_kind in {"artifact", "attachment"} or topic.get("publish") is False:
                    continue
                topic_uid = topic_public_uid(topic)
                content_path = source_dir / str(topic["content_ref"])
                section_row["topics"].append(
                    {
                        "uid": topic_uid,
                        "title": str(topic["title"]),
                        "fragments": markdown_navigation_headings(content_path, topic_uid, str(topic["title"])),
                    }
                )
        outline.append(section_row)
    return outline


def html_toc(outline: list[dict[str, Any]], book_uid: str, book_title: str) -> str:
    expanded = " open" if len(outline) == 1 else ""
    parts = [
        "<aside id='std_toc' class='toc-panel machine-anchor'>",
        "<nav aria-label='Оглавление книги'>",
        "<h2>Оглавление</h2>",
        f"<a class='toc-section-start' href='#{html.escape(book_uid)}'>{html.escape(book_title)}</a>",
        "<ol class='toc-sections'>",
    ]
    for section in outline:
        parts.append(f"<li><details{expanded}><summary>{html.escape(section['title'])}</summary>")
        parts.append(f"<a class='toc-section-start' href='#{html.escape(section['uid'])}'>К началу раздела</a>")
        if section["topics"]:
            parts.append("<ol class='toc-topics'>")
            for topic in section["topics"]:
                parts.append(f"<li><a href='#{html.escape(topic['uid'])}'>{html.escape(topic['title'])}</a>")
                if topic["fragments"]:
                    parts.append("<ol class='toc-fragments'>")
                    parts.extend(
                        f"<li><a href='#{html.escape(fragment['uid'])}'>{html.escape(fragment['title'])}</a></li>"
                        for fragment in topic["fragments"]
                    )
                    parts.append("</ol>")
                parts.append("</li>")
            parts.append("</ol>")
        parts.append("</details></li>")
    parts.extend(["</ol>", "</nav>", "</aside>"])
    return "\n".join(parts)


def topic_navigation_lookup(outline: list[dict[str, Any]]) -> dict[str, dict[str, dict[str, str] | None]]:
    topics = [topic for section in outline for topic in section["topics"]]
    result: dict[str, dict[str, dict[str, str] | None]] = {}
    for index, topic in enumerate(topics):
        result[topic["uid"]] = {
            "previous": topics[index - 1] if index else None,
            "next": topics[index + 1] if index + 1 < len(topics) else None,
        }
    return result


def html_topic_navigation(topic_uid: str, lookup: dict[str, dict[str, dict[str, str] | None]]) -> str:
    neighbors = lookup.get(topic_uid, {})
    previous = neighbors.get("previous")
    following = neighbors.get("next")
    previous_html = (
        f"<a rel='prev' href='#{html.escape(previous['uid'])}'>← {html.escape(previous['title'])}</a>"
        if previous
        else "<span></span>"
    )
    next_html = (
        f"<a rel='next' href='#{html.escape(following['uid'])}'>{html.escape(following['title'])} →</a>"
        if following
        else "<span></span>"
    )
    return f"<nav class='topic-navigation' aria-label='Навигация по темам'>{previous_html}<a href='#std_toc'>Оглавление</a>{next_html}</nav>"


def docx_bookmark_name(uid: str) -> str:
    bookmark_name = re.sub(r"[^A-Za-z0-9_]", "_", uid)
    if not bookmark_name or not bookmark_name[0].isalpha():
        bookmark_name = "uid_" + bookmark_name
    if len(bookmark_name) > 40:
        bookmark_name = f"{bookmark_name[:27]}_{hashlib.sha1(uid.encode('utf-8')).hexdigest()[:12]}"
    return bookmark_name


def add_docx_bookmark(paragraph: Any, uid: str, bookmark_id: int) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    bookmark_name = docx_bookmark_name(uid)
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), bookmark_name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_docx_internal_link(paragraph: Any, text: str, target_uid: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), docx_bookmark_name(target_uid))
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    run.append(properties)
    label = OxmlElement("w:t")
    label.text = text
    run.append(label)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_docx_external_link(paragraph: Any, text: str, url: str, segment: dict[str, Any]) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE

    relationship_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    if segment.get("bold"):
        properties.append(OxmlElement("w:b"))
    if segment.get("italic"):
        properties.append(OxmlElement("w:i"))
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    properties.append(color)
    run.append(properties)
    label = OxmlElement("w:t")
    label.text = text
    run.append(label)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_docx_styled_text(paragraph: Any, text: str, hint: dict[str, Any] | None) -> None:
    if not hint or parity_plain_text(hint.get("text") or "") != parity_plain_text(markdown_plain_text(text)):
        paragraph.add_run(text)
        return
    for segment in hint.get("segments") or []:
        segment_text = str(segment.get("text") or "")
        link = str(segment.get("link") or "")
        if link.startswith(("http://", "https://", "mailto:")):
            add_docx_external_link(paragraph, segment_text, link, segment)
            continue
        run = paragraph.add_run(segment_text)
        run.bold = bool(segment.get("bold"))
        run.italic = bool(segment.get("italic"))
        run.underline = bool(segment.get("underline"))


def add_docx_styled_paragraph(doc: Any, text: str, hint: dict[str, Any] | None, style: str | None = None) -> Any:
    paragraph = doc.add_paragraph(style=style)
    add_docx_styled_text(paragraph, text, hint)
    return paragraph


def add_docx_picture_resilient(owner: Any, image_path: Path, width: Any) -> None:
    """Insert an image, normalizing valid JPEGs rejected by python-docx."""
    try:
        owner.add_picture(str(image_path), width=width)
        return
    except Exception as original_exc:  # noqa: BLE001
        try:
            from PIL import Image

            with tempfile.TemporaryDirectory(prefix="std-docx-image-") as raw:
                normalized = Path(raw) / "normalized.png"
                with Image.open(image_path) as image:
                    if image.mode not in {"RGB", "RGBA"}:
                        image = image.convert("RGB")
                    image.save(normalized, format="PNG")
                owner.add_picture(str(normalized), width=width)
                return
        except Exception:  # noqa: BLE001
            raise original_exc


def render_html(source_dir: Path, output: Path, profile: str, section_uids: list[str] | None = None) -> None:
    book, sections, topics, _ = collect_units(source_dir)
    topics_by_uid = {row["uid"]: row for row in topics}
    section_by_uid = {row["uid"]: row for row in sections}
    table_layout_contract = load_table_layout_contract()
    positioned_assets = positioned_assets_by_tab(source_dir)
    book_uid = str(book.get("semantic_uid") or book.get("uid") or "std_iridi")
    outline = navigation_outline(source_dir, book, sections, topics, profile, section_uids)
    topic_navigation = topic_navigation_lookup(outline)
    body: list[str] = [f"<h1 id='{html.escape(book_uid)}' class='machine-anchor' data-book-uid='{html.escape(book_uid)}'>{html.escape(book['title'])}</h1>"]
    for ref in selected_section_refs(book, sections, section_uids):
        section = section_by_uid[ref]
        topic_style_hints = section_topic_style_hints(source_dir, section)
        section_uid = str(section.get("semantic_uid") or section["uid"])
        body.append(
            f"<h1 id='{html.escape(section_uid)}' class='machine-anchor' data-section-uid='{html.escape(section_uid)}'>"
            f"{html.escape(str(section.get('display_number') or ''))} {html.escape(section['title'])}</h1>"
        )
        if profile == "legacy-fidelity":
            body.extend(markdown_html(source_dir / "sections" / ref / "legacy.md"))
        else:
            for topic in ordered_section_topics(section, topics_by_uid):
                node_kind = str(topic.get("node_kind") or "content")
                if node_kind == "artifact" or topic.get("publish") is False:
                    continue
                public_uid = topic_public_uid(topic)
                if node_kind != "attachment":
                    body.append(html_topic_navigation(public_uid, topic_navigation))
                    body.append(
                        f"<h2 id='{html.escape(public_uid)}' class='machine-anchor topic-{html.escape(node_kind)}' "
                        f"data-topic-uid='{html.escape(public_uid)}' data-legacy-uid='{html.escape(str(topic['uid']))}' "
                        f"data-node-kind='{html.escape(node_kind)}'>{html.escape(topic['title'])}</h2>"
                    )
                body.extend(
                    markdown_html(
                        source_dir / topic["content_ref"],
                        skip_initial_heading="" if node_kind == "attachment" else str(topic["title"]),
                        table_layouts=table_layout_contract.get(str(topic["content_ref"])),
                        suppress_initial_heading=node_kind == "attachment",
                        fragment_uid_prefix=f"std_fragment_{public_uid.removeprefix('std_topic_')}",
                        style_hints=topic_style_hints.get(str(topic["uid"])),
                    )
                )
                if node_kind == "gap":
                    body.append("<aside class='gap-note' data-coverage-status='gap'>Раздел присутствует в исходной книге, но пока не наполнен.</aside>")
        unplaced = positioned_assets.get(str(section.get("source_tab_id")), [])
        if unplaced:
            body.append("<aside class='unplaced-assets'><strong>Неразмещенные иллюстрации исходника</strong><br>Google Docs не передал надежную привязку этих плавающих объектов к абзацу; изображения сохранены и требуют ручной проверки места.</aside>")
            for asset in unplaced:
                path = source_dir / str(asset.get("path") or "")
                if path.is_file():
                    body.append(
                        f"<figure data-asset-uid='{html.escape(str(asset.get('object_id')))}' data-placement-status='needs_manual_placement'>"
                        f"<img src='assets/{html.escape(path.name)}' alt='{html.escape(str(asset.get('object_id')))}'>"
                        f"<figcaption>{html.escape(str(asset.get('object_id')))} — исходное плавающее изображение, место требует проверки.</figcaption></figure>"
                    )
    page = (
        "<!doctype html><html lang='ru'><head><meta charset='utf-8'>"
        "<meta name='viewport' content='width=device-width,initial-scale=1'>"
        f"<title>{html.escape(str(book['title']))}</title><style>{HTML_CSS}</style></head><body>"
        "<div class='book-shell'>"
        + html_toc(outline, book_uid, str(book["title"]))
        + "<main class='book-content'>"
        + "\n".join(body)
        + "</main></div><a class='back-to-toc' href='#std_toc' aria-label='Вернуться к оглавлению'>Оглавление ↑</a>"
        "</body></html>\n"
    )
    atomic_write_text(output, page)


def render_docx(source_dir: Path, output: Path, profile: str, section_uids: list[str] | None = None) -> None:
    try:
        from docx import Document
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.shared import Mm, Pt
    except ImportError as exc:  # pragma: no cover
        raise SystemExit("python-docx is required for DOCX build") from exc
    book, sections, topics, _ = collect_units(source_dir)
    doc = Document()
    fixed_time = datetime(2000, 1, 1, tzinfo=timezone.utc)
    doc.core_properties.created = fixed_time
    doc.core_properties.modified = fixed_time
    doc.core_properties.revision = 1
    sec = doc.sections[0]
    sec.top_margin, sec.right_margin, sec.bottom_margin, sec.left_margin = Mm(18), Mm(18), Mm(20), Mm(20)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    bookmark_counter = 1
    title_paragraph = doc.add_heading(book["title"], 0)
    book_uid = str(book.get("semantic_uid") or book.get("uid") or "std_iridi")
    add_docx_bookmark(title_paragraph, book_uid, bookmark_counter)
    bookmark_counter += 1
    if profile == "standard-normalized":
        doc.add_paragraph("Нормализованное представление: единая структура разделов и адресуемых тем.")
    section_by_uid = {row["uid"]: row for row in sections}
    topics_by_uid = {row["uid"]: row for row in topics}
    outline = navigation_outline(source_dir, book, sections, topics, profile, section_uids)
    topic_navigation = topic_navigation_lookup(outline)
    toc_heading = doc.add_heading("Оглавление", 1)
    add_docx_bookmark(toc_heading, "std_toc", bookmark_counter)
    bookmark_counter += 1
    book_link = doc.add_paragraph()
    add_docx_internal_link(book_link, str(book["title"]), book_uid)
    for section_row in outline:
        section_link = doc.add_paragraph()
        add_docx_internal_link(section_link, section_row["title"], section_row["uid"])
        for topic_row in section_row["topics"]:
            topic_link = doc.add_paragraph()
            topic_link.paragraph_format.left_indent = Mm(6)
            add_docx_internal_link(topic_link, topic_row["title"], topic_row["uid"])
            for fragment in topic_row["fragments"]:
                fragment_link = doc.add_paragraph()
                fragment_link.paragraph_format.left_indent = Mm(12)
                add_docx_internal_link(fragment_link, fragment["title"], fragment["uid"])
    doc.add_page_break()
    table_layout_contract = load_table_layout_contract()
    positioned_assets = positioned_assets_by_tab(source_dir)
    for section_index, ref in enumerate(selected_section_refs(book, sections, section_uids)):
        section = section_by_uid[ref]
        topic_style_hints = section_topic_style_hints(source_dir, section)
        if profile == "standard-normalized" and section_index:
            doc.add_page_break()
        section_heading = doc.add_heading(f"{section.get('display_number') or ''} {section['title']}".strip(), 1)
        add_docx_bookmark(section_heading, str(section.get("semantic_uid") or section["uid"]), bookmark_counter)
        bookmark_counter += 1
        section_backlink = doc.add_paragraph()
        add_docx_internal_link(section_backlink, "↑ Оглавление", "std_toc")
        entries: list[tuple[Path, str | None, dict[int, dict[str, Any]] | None, dict[str, Any] | None]]
        if profile == "legacy-fidelity":
            entries = [(source_dir / "sections" / ref / "legacy.md", None, None, None)]
        else:
            entries = [
                (
                    source_dir / topic["content_ref"],
                    None if topic.get("node_kind") == "attachment" else str(topic["title"]),
                    table_layout_contract.get(str(topic["content_ref"])),
                    topic,
                )
                for topic in ordered_section_topics(section, topics_by_uid)
                if topic.get("node_kind") != "artifact" and topic.get("publish") is not False
            ]
        for path, normalized_topic_title, table_layouts, topic_meta in entries:
            normalized_entry = profile == "standard-normalized"
            node_kind = str((topic_meta or {}).get("node_kind") or "content")
            public_uid = topic_public_uid(topic_meta) if topic_meta else None
            style_hints = topic_style_hints.get(str((topic_meta or {}).get("uid") or ""))
            if normalized_topic_title:
                neighbors = topic_navigation.get(public_uid or "", {})
                navigation_paragraph = doc.add_paragraph()
                if previous := neighbors.get("previous"):
                    add_docx_internal_link(navigation_paragraph, f"← {previous['title']}", previous["uid"])
                    navigation_paragraph.add_run("   |   ")
                add_docx_internal_link(navigation_paragraph, "Оглавление", "std_toc")
                if following := neighbors.get("next"):
                    navigation_paragraph.add_run("   |   ")
                    add_docx_internal_link(navigation_paragraph, f"{following['title']} →", following["uid"])
                paragraph = doc.add_heading(normalized_topic_title, 2)
                if public_uid:
                    add_docx_bookmark(paragraph, public_uid, bookmark_counter)
                    bookmark_counter += 1
            lines = logical_markdown_lines(path.read_text(encoding="utf-8"))
            index = 0
            table_index = 0
            checked_initial_heading = False
            heading_occurrences: dict[str, int] = defaultdict(int)
            while index < len(lines):
                line = lines[index]
                stripped = line.strip()
                if not stripped:
                    index += 1
                    continue
                if not checked_initial_heading:
                    checked_initial_heading = True
                    heading = markdown_heading(stripped)
                    if node_kind == "attachment" and heading:
                        stripped = heading[1]
                    elif normalized_topic_title and heading and heading[1] == normalized_topic_title:
                        take_style_hint(style_hints, heading[1])
                        index += 1
                        continue
                if stripped.startswith("| "):
                    rows: list[list[str]] = []
                    while index < len(lines) and lines[index].strip().startswith("|"):
                        cells = markdown_table_cells(lines[index])
                        if not markdown_is_separator(cells):
                            rows.append(cells)
                        index += 1
                    if rows:
                        if normalized_entry:
                            rows, header_rows, column_widths = analyze_table_rows(rows, (table_layouts or {}).get(table_index))
                        else:
                            rows, header_rows, column_widths = legacy_table_rows(rows)
                        table_index += 1
                        width = len(rows[0])
                        table = doc.add_table(rows=len(rows), cols=width)
                        table.style = "Table Grid"
                        set_docx_table_geometry(table, column_widths)
                        for row_index, row in enumerate(rows):
                            for col_index in range(width):
                                cell = table.cell(row_index, col_index)
                                column_width_mm = 160 * column_widths[col_index] / 100.0
                                cell.width = Mm(column_width_mm)
                                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                                value = row[col_index] if col_index < len(row) else ""
                                image_matches = list(re.finditer(r"!\[([^]]*)\]\(([^)]+)\)", value))
                                if image_matches:
                                    paragraph = cell.paragraphs[0]
                                    paragraph.clear()
                                    plain = re.sub(r"!\[[^]]*\]\([^)]+\)", "", value).strip()
                                    if plain:
                                        paragraph.add_run(plain)
                                    for image_match in image_matches:
                                        image_path = (path.parent / image_match.group(2)).resolve()
                                        if image_path.is_file():
                                            try:
                                                image_width_mm = max(20.0, min(70.0, column_width_mm - 6.0))
                                                add_docx_picture_resilient(paragraph.add_run(), image_path, Mm(image_width_mm))
                                            except Exception:  # noqa: BLE001
                                                paragraph.add_run(f" [Изображение: {image_path.name}]")
                                else:
                                    cell.text = value
                                if row_index < header_rows:
                                    for run in cell.paragraphs[0].runs:
                                        run.bold = True
                    continue
                source_heading = markdown_heading(stripped)
                renderable_line = source_heading[1] if source_heading else stripped
                image_matches = list(re.finditer(r"!\[([^]]*)\]\(([^)]+)\)", renderable_line))
                if image_matches:
                    hint = take_style_hint(style_hints, renderable_line)
                    plain = re.sub(r"!\[[^]]*\]\([^)]+\)", "", renderable_line).strip()
                    if source_heading and plain:
                        doc.add_heading(plain, rendered_heading_level(source_heading[0], normalized=normalized_entry))

                    def add_plain_piece(value: str) -> None:
                        value = value.strip()
                        if not value or source_heading:
                            return
                        piece_hint = hint if parity_plain_text(value.lstrip("- ")) == parity_plain_text((hint or {}).get("text") or "") else None
                        if value.startswith("- "):
                            add_docx_styled_paragraph(doc, value[2:], piece_hint, style="List Bullet")
                        else:
                            add_docx_styled_paragraph(doc, value, piece_hint)

                    cursor = 0
                    for image_match in image_matches:
                        add_plain_piece(renderable_line[cursor:image_match.start()])
                        image_path = (path.parent / image_match.group(2)).resolve()
                        if image_path.is_file():
                            try:
                                add_docx_picture_resilient(doc, image_path, Mm(160))
                            except Exception:  # noqa: BLE001
                                doc.add_paragraph(f"[Изображение: {image_path.name}]")
                        else:
                            doc.add_paragraph(image_match.group(0))
                        cursor = image_match.end()
                    add_plain_piece(renderable_line[cursor:])
                elif heading := markdown_heading(stripped):
                    take_style_hint(style_hints, heading[1])
                    paragraph = doc.add_heading(heading[1], rendered_heading_level(heading[0], normalized=normalized_entry))
                    if public_uid:
                        heading_occurrences[heading[1]] += 1
                        fragment_uid = fragment_anchor_uid(public_uid, heading[1], heading_occurrences[heading[1]])
                        add_docx_bookmark(paragraph, fragment_uid, bookmark_counter)
                        bookmark_counter += 1
                elif stripped.startswith("- "):
                    add_docx_styled_paragraph(doc, stripped[2:], take_style_hint(style_hints, stripped[2:]), style="List Bullet")
                elif stripped.startswith("| "):
                    doc.add_paragraph(stripped)
                else:
                    add_docx_styled_paragraph(doc, stripped, take_style_hint(style_hints, stripped))
                index += 1
            if node_kind == "gap":
                doc.add_paragraph("Раздел присутствует в исходной книге, но пока не наполнен.")
        unplaced = positioned_assets.get(str(section.get("source_tab_id")), [])
        if unplaced:
            doc.add_heading("Неразмещенные иллюстрации исходника", 2)
            doc.add_paragraph(
                "Google Docs не передал надежную привязку этих плавающих объектов к абзацу. "
                "Изображения сохранены; место требует ручной проверки по исходной книге."
            )
            for asset in unplaced:
                image_path = source_dir / str(asset.get("path") or "")
                if image_path.is_file():
                    try:
                        add_docx_picture_resilient(doc, image_path, Mm(100))
                    except Exception:  # noqa: BLE001
                        doc.add_paragraph(f"[Изображение: {image_path.name}]")
                    doc.add_paragraph(f"{asset.get('object_id')} — исходное плавающее изображение; место требует проверки.")
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    normalize_zip_archive(output)


def build_cmd(args: argparse.Namespace) -> int:
    source_dir = Path(args.source_dir).resolve()
    output_dir = Path(args.output_dir).resolve()
    requested_sections = list(getattr(args, "sections", None) or [])
    book, sections, topics, _ = collect_units(source_dir)
    section_refs = selected_section_refs(book, sections, requested_sections)
    section_by_uid = {row["uid"]: row for row in sections}
    topics_by_uid = {row["uid"]: row for row in topics}
    topic_refs = [
        topic["uid"]
        for ref in section_refs
        for topic in ordered_section_topics(section_by_uid[ref], topics_by_uid)
    ]
    output_dir.mkdir(parents=True, exist_ok=True)
    source_assets = source_dir / "assets"
    if source_assets.is_dir():
        shutil.copytree(source_assets, output_dir / "assets", dirs_exist_ok=True)
    if args.format in {"html", "all"}:
        render_html(source_dir, output_dir / f"standard-{args.profile}.html", args.profile, section_refs)
    if args.format in {"docx", "all"}:
        render_docx(source_dir, output_dir / f"standard-{args.profile}.docx", args.profile, section_refs)
    manifest = {
        "profile": args.profile,
        "built_at": utc_now(),
        "source_book_digest": sha256_file(source_dir / "book.yaml"),
        "scope": "full_book" if not requested_sections else "selected_sections",
        "section_refs": section_refs,
        "counts": {"sections": len(section_refs), "topics": len(topic_refs)},
        "outputs": [
            {"path": path.name, "sha256": sha256_file(path), "bytes": path.stat().st_size}
            for path in sorted(output_dir.glob(f"standard-{args.profile}.*"))
        ],
    }
    write_json(output_dir / f"build-{args.profile}.json", manifest)
    print(json.dumps(manifest, ensure_ascii=False))
    return 0


def canonical_topic_rows(book: dict[str, Any], sections: list[dict[str, Any]], topics: list[dict[str, Any]]) -> list[dict[str, Any]]:
    section_by_uid = {row["uid"]: row for row in sections}
    topic_by_uid = {row["uid"]: row for row in topics}
    return [
        topic_by_uid[topic_uid]
        for section_uid in book.get("section_refs") or []
        for topic_uid in section_by_uid.get(section_uid, {}).get("topic_refs") or []
        if topic_uid in topic_by_uid
    ]


def attachment_payload(content: str) -> str:
    lines = logical_markdown_lines(content)
    if lines:
        heading = markdown_heading(lines[0].strip())
        if heading:
            lines[0] = heading[1]
    return "\n".join(lines).strip()


def rewrite_package_asset_refs(content: str, source_content: Path) -> tuple[str, list[Path]]:
    assets: list[Path] = []

    def replace(match: re.Match[str]) -> str:
        image_path = (source_content.parent / match.group(2)).resolve()
        if image_path.is_file():
            assets.append(image_path)
            return f"{match.group(1)}(../../assets/{image_path.name})"
        return match.group(0)

    return re.sub(r"(!\[[^]]*\])\(([^)]+)\)", replace, content), assets


def markdown_fragments(topic_uid: str, section_uid: str, content: str, content_ref: str) -> list[dict[str, Any]]:
    lines = logical_markdown_lines(content)
    fragments: list[dict[str, Any]] = []
    current_heading = ""
    buffer: list[str] = []
    occurrence: dict[str, int] = defaultdict(int)

    def flush() -> None:
        nonlocal buffer
        raw = "\n".join(buffer).strip()
        buffer = []
        plain = markdown_plain_text(raw)
        if not raw or not plain:
            return
        digest = hashlib.sha1(f"{topic_uid}|{current_heading}|{plain}".encode("utf-8")).hexdigest()[:12]
        occurrence[digest] += 1
        fragment_uid = f"std_fragment_{digest}_{occurrence[digest]}"
        fragments.append(
            {
                "uid": fragment_uid,
                "type": "fragment",
                "topic_uid": topic_uid,
                "section_uid": section_uid,
                "heading": current_heading,
                "ordinal": len(fragments) + 1,
                "text": raw,
                "plain_text": plain,
                "content_ref": content_ref,
            }
        )

    for line in lines:
        heading = markdown_heading(line.strip())
        if heading:
            flush()
            current_heading = heading[1]
            continue
        if not line.strip():
            flush()
            continue
        buffer.append(line)
    flush()
    if not fragments and markdown_plain_text(content):
        plain = markdown_plain_text(content)
        digest = hashlib.sha1(f"{topic_uid}|{plain}".encode("utf-8")).hexdigest()[:12]
        fragments.append(
            {
                "uid": f"std_fragment_{digest}_1",
                "type": "fragment",
                "topic_uid": topic_uid,
                "section_uid": section_uid,
                "heading": "",
                "ordinal": 1,
                "text": content.strip(),
                "plain_text": plain,
                "content_ref": content_ref,
            }
        )
    return fragments


def index_cmd(args: argparse.Namespace) -> int:
    source_dir = Path(args.source_dir).resolve()
    output_dir = Path(args.output_dir).resolve()
    if output_dir.exists() and any(output_dir.iterdir()):
        if not (output_dir / "package.yaml").exists() and not args.force:
            raise SystemExit(f"Refusing to replace a non-package directory: {output_dir}. Use --force only for a deliberate generated target.")
        for name in ("content", "indexes", "assets"):
            target = (output_dir / name).resolve()
            if target.parent != output_dir:
                raise SystemExit(f"Unsafe package reset refused: {target}")
            if target.exists():
                shutil.rmtree(target)
        for name in (
            "topics.jsonl", "fragments.jsonl", "rules.jsonl", "entities.jsonl", "relations.jsonl", "unplaced_assets.jsonl",
            "aliases.json", "changes.json", "navigation.json", "package.yaml", "START_HERE.md",
        ):
            target = output_dir / name
            if target.exists():
                target.unlink()
    output_dir.mkdir(parents=True, exist_ok=True)
    content_dir = output_dir / "content" / "by-uid"
    content_dir.mkdir(parents=True, exist_ok=True)
    book, sections, topics, rules = collect_units(source_dir)
    section_by_uid = {row["uid"]: row for row in sections}
    ordered_topics = canonical_topic_rows(book, sections, topics)
    semantic_by_storage = {str(row["uid"]): topic_public_uid(row) for row in ordered_topics}
    primary_topics = [
        row for row in ordered_topics
        if row.get("node_kind") not in {"artifact", "attachment"} and row.get("publish") is not False
    ]
    attachments_by_target: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for topic in ordered_topics:
        if topic.get("node_kind") == "attachment" and topic.get("attached_to_uid"):
            attachments_by_target[str(topic["attached_to_uid"])].append(topic)
    topic_rows: list[dict[str, Any]] = []
    fragment_rows: list[dict[str, Any]] = []
    relation_rows: list[dict[str, Any]] = []
    required_assets: dict[str, Path] = {}
    indexes: dict[str, dict[str, list[str]]] = {
        name: defaultdict(list) for name in ["audience", "job", "domain", "lifecycle", "entity", "node_kind", "alias"]
    }
    for topic in primary_topics:
        row = dict(topic)
        source_content = source_dir / row["content_ref"]
        semantic_uid = topic_public_uid(row)
        target_content = content_dir / f"{semantic_uid}.md"
        content_text, asset_paths = rewrite_package_asset_refs(source_content.read_text(encoding="utf-8"), source_content)
        for attachment in attachments_by_target.get(semantic_uid, []):
            attachment_path = source_dir / attachment["content_ref"]
            payload, attachment_assets = rewrite_package_asset_refs(
                attachment_payload(attachment_path.read_text(encoding="utf-8")), attachment_path
            )
            if payload:
                content_text = content_text.rstrip() + "\n\n" + payload + "\n"
            asset_paths.extend(attachment_assets)
            relation_rows.append(
                {
                    "uid": stable_uid("std_relation", topic_public_uid(attachment), semantic_uid, "attached_to"),
                    "type": "relation",
                    "relation_type": "attached_to",
                    "from_uid": topic_public_uid(attachment),
                    "to_uid": semantic_uid,
                    "source_ref": attachment["content_ref"],
                }
            )
        for asset_path in asset_paths:
            required_assets[asset_path.name] = asset_path
        atomic_write_text(target_content, content_text)
        storage_uid = str(row["uid"])
        row["storage_uid"] = storage_uid
        row["uid"] = semantic_uid
        row["content_ref"] = target_content.relative_to(output_dir).as_posix()
        row["section_title"] = section_by_uid.get(row.get("parent_uid"), {}).get("title")
        row["semantic_parent_uid"] = str(row.get("semantic_parent_uid") or row.get("parent_uid"))
        row["digest"] = sha256_file(target_content)
        topic_rows.append(row)
        for field, index_name in [("audiences", "audience"), ("jobs", "job"), ("domains", "domain"), ("lifecycle", "lifecycle"), ("entity_refs", "entity")]:
            for value in row.get(field) or []:
                indexes[index_name][str(value)].append(semantic_uid)
        indexes["node_kind"][str(row.get("node_kind") or "content")].append(semantic_uid)
        for alias in row.get("aliases") or []:
            indexes["alias"][normalize_search_text(str(alias))].append(semantic_uid)
        parent_uid = str(row.get("semantic_parent_uid") or "")
        if parent_uid and parent_uid != row.get("parent_uid"):
            relation_rows.append(
                {
                    "uid": stable_uid("std_relation", semantic_uid, parent_uid, "child_of"),
                    "type": "relation",
                    "relation_type": "child_of",
                    "from_uid": semantic_uid,
                    "to_uid": parent_uid,
                }
            )
        fragments = markdown_fragments(semantic_uid, str(row.get("parent_uid")), content_text, row["content_ref"])
        fragment_rows.extend(fragments)

    positioned = positioned_assets_by_tab(source_dir)
    unplaced_rows: list[dict[str, Any]] = []
    for section in sections:
        if section["uid"] not in (book.get("section_refs") or []):
            continue
        for asset in positioned.get(str(section.get("source_tab_id")), []):
            source_asset = source_dir / str(asset.get("path") or "")
            if source_asset.is_file():
                required_assets[source_asset.name] = source_asset
            asset_uid = str(asset.get("object_id"))
            unplaced_rows.append(
                {
                    "uid": asset_uid,
                    "type": "asset",
                    "asset_type": "image",
                    "path": f"assets/{source_asset.name}",
                    "checksum": asset.get("sha256"),
                    "section_uid": section["uid"],
                    "placement_status": "needs_manual_placement",
                    "source_tab_id": section.get("source_tab_id"),
                }
            )
            relation_rows.append(
                {
                    "uid": stable_uid("std_relation", section["uid"], asset_uid, "has_unplaced_asset"),
                    "type": "relation",
                    "relation_type": "has_unplaced_asset",
                    "from_uid": section["uid"],
                    "to_uid": asset_uid,
                }
            )

    for name, source_asset in sorted(required_assets.items()):
        target = output_dir / "assets" / name
        target.parent.mkdir(parents=True, exist_ok=True)
        if not target.exists() or sha256_file(target) != sha256_file(source_asset):
            shutil.copy2(source_asset, target)

    semantic_rules: list[dict[str, Any]] = []
    for original in rules:
        rule = dict(original)
        if rule.get("parent_uid") in semantic_by_storage:
            rule["parent_uid"] = semantic_by_storage[str(rule["parent_uid"])]
        semantic_rules.append(rule)
    for filename, rows in [
        ("topics.jsonl", topic_rows),
        ("fragments.jsonl", fragment_rows),
        ("rules.jsonl", semantic_rules),
        ("entities.jsonl", []),
        ("relations.jsonl", relation_rows),
        ("unplaced_assets.jsonl", unplaced_rows),
    ]:
        atomic_write_text(output_dir / filename, "".join(json.dumps(row, ensure_ascii=False, sort_keys=True) + "\n" for row in rows))
    write_json(output_dir / "aliases.json", {row["uid"]: row.get("aliases") or [] for row in topic_rows})
    write_json(output_dir / "changes.json", [])
    navigation_topics = [
        {
            "uid": row["uid"],
            "storage_uid": row.get("storage_uid"),
            "title": row.get("title"),
            "display_number": row.get("display_number"),
            "node_kind": row.get("node_kind"),
            "parent_uid": row.get("semantic_parent_uid"),
            "coverage_status": row.get("coverage_status"),
        }
        for row in topic_rows
    ]
    navigation_sections = []
    published_topic_uids = {str(row["uid"]) for row in topic_rows}
    for section in sections:
        nav_section = dict(section)
        nav_section["topic_refs"] = [
            semantic_by_storage[storage_uid]
            for storage_uid in section.get("topic_refs") or []
            if semantic_by_storage.get(storage_uid) in published_topic_uids
        ]
        navigation_sections.append(nav_section)
    write_json(output_dir / "navigation.json", {"book": book, "sections": navigation_sections, "topics": navigation_topics})
    for name, values in indexes.items():
        write_json(output_dir / "indexes" / f"by_{name}.json", dict(sorted(values.items())))
    package = {
        "schema_version": "1.0",
        "package_uid": stable_uid("std_pkg", str(book.get("baseline_uid")), args.release),
        "release": args.release,
        "generated_at": utc_now(),
        "source_book_digest": sha256_file(source_dir / "book.yaml"),
        "counts": {
            "sections": len(sections),
            "topics": len(topic_rows),
            "fragments": len(fragment_rows),
            "rules": len(semantic_rules),
            "relations": len(relation_rows),
            "assets": len(required_assets),
            "unplaced_assets": len(unplaced_rows),
        },
        "drafts_included": False,
        "private_sources_included": False,
        "buffer_included": False,
        "query_contract": "facets_then_weighted_semantic_recall_then_uid_content",
    }
    write_yaml(output_dir / "package.yaml", package)
    atomic_write_text(
        output_dir / "START_HERE.md",
        "# Standard knowledge package\n\n"
        f"Release: `{args.release}`. Read `package.yaml`, then route by audience/job/domain/alias indexes. "
        "Use `topics.jsonl`, `fragments.jsonl`, `relations.jsonl` and `rules.jsonl`. "
        "Only content/gap nodes are queryable. Cite release, semantic topic UID, fragment UID and content_ref. "
        "Never infer a documented rule from absence; Buffer, artifacts and unresolved placements are excluded from answers.\n",
    )
    print(json.dumps(package, ensure_ascii=False))
    return 0


def propose_rules_cmd(args: argparse.Namespace) -> int:
    source_dir = Path(args.source_dir).resolve()
    section_dir = source_dir / "sections" / args.section
    section = read_yaml(section_dir / "section.yaml")
    patterns = [
        (re.compile(r"\b(не допускается|нельзя|не следует|не рекомендуется)\b", re.I), "prohibition", "must_not"),
        (re.compile(r"\b(должен|должна|должны|должно|необходимо|обязательно)\b", re.I), "requirement", "must"),
        (re.compile(r"\b(следует|рекомендуется|рекомендуются|рекомендуемое|рекомендуется)\b", re.I), "recommendation", "should"),
    ]
    candidates: list[dict[str, Any]] = []
    seen: set[str] = set()
    topic_paths = {path.parent.name: path for path in section_dir.glob("topics/*/content.md")}
    for topic_uid in section.get("topic_refs") or []:
        path = topic_paths.get(topic_uid)
        if not path:
            continue
        text = path.read_text(encoding="utf-8")
        for line_number, line in enumerate(logical_markdown_lines(text), start=1):
            plain = re.sub(r"^\s*(?:[-#]+|\|)\s*", "", line).strip(" |").strip()
            if not plain:
                continue
            for statement in re.split(r"(?<=[.!?])\s+", plain):
                normalized = re.sub(r"\s+", " ", statement).strip()
                if len(normalized) < 20:
                    continue
                for pattern, claim_type, normative_level in patterns:
                    match = pattern.search(normalized)
                    if not match:
                        continue
                    key = normalized.casefold()
                    if key in seen:
                        break
                    seen.add(key)
                    candidates.append(
                        {
                            "candidate_uid": stable_uid("rule_candidate", topic_uid, str(line_number), normalized),
                            "status": "needs_expert_review",
                            "parent_uid": topic_uid,
                            "claim_type_suggestion": claim_type,
                            "normative_level_suggestion": normative_level,
                            "statement": normalized,
                            "trigger": match.group(1),
                            "source_ref": f"{path.relative_to(source_dir).as_posix()}#L{line_number}",
                            "confidence": 0.55,
                            "warning": "Lexical candidate only; context, scope, applicability, exceptions and authority are not approved.",
                        }
                    )
                    break
    report = {
        "schema_version": "1.0",
        "section_uid": section["uid"],
        "generated_at": utc_now(),
        "status": "review_required",
        "candidate_count": len(candidates),
        "approved_rule_count": 0,
        "candidates": candidates,
    }
    write_yaml(Path(args.output).resolve(), report)
    print(json.dumps({key: report[key] for key in ("section_uid", "status", "candidate_count", "approved_rule_count")}, ensure_ascii=False))
    return 0


def normalize_search_text(text: str) -> str:
    value = unicodedata.normalize("NFKC", text).casefold().replace("ё", "е")
    value = re.sub(r"[^\w]+", " ", value, flags=re.UNICODE)
    return re.sub(r"\s+", " ", value).strip()


RUSSIAN_ENDINGS = tuple(
    sorted(
        {
            "иями", "ями", "ами", "ией", "иям", "ием", "иях", "ость", "ости", "остью", "ение", "ения", "ений",
            "ого", "ему", "ому", "ими", "ыми", "его", "ая", "яя", "ое", "ее", "ие", "ые", "ой", "ий", "ый",
            "ую", "юю", "ах", "ях", "ам", "ям", "ом", "ем", "ов", "ев", "ей", "ия", "ие", "ий", "ы", "и",
            "а", "я", "у", "ю", "е", "о",
        },
        key=len,
        reverse=True,
    )
)


def search_token_variants(token: str) -> set[str]:
    value = token.casefold().replace("ё", "е")
    variants = {value}
    if len(value) >= 6:
        variants.add("^" + value[:6])
    if re.search(r"[а-я]", value) and len(value) >= 5:
        for ending in RUSSIAN_ENDINGS:
            if value.endswith(ending) and len(value) - len(ending) >= 4:
                variants.add(value[: -len(ending)])
                break
    return variants


def token_set(text: str) -> set[str]:
    result: set[str] = set()
    for token in TOKEN_RE.findall(normalize_search_text(text)):
        if token in STOPWORDS:
            continue
        result.update(search_token_variants(token))
    return result


def concept_tokens(text: str) -> set[str]:
    return {
        token.casefold().replace("ё", "е")
        for token in TOKEN_RE.findall(normalize_search_text(text))
        if token.casefold().replace("ё", "е") not in STOPWORDS
    }


def matched_query_concepts(query: set[str], document_text: str) -> set[str]:
    document = concept_tokens(document_text)
    matched: set[str] = set()
    for query_token in query:
        for document_token in document:
            if query_token == document_token or (
                len(query_token) >= 6 and len(document_token) >= 6 and query_token[:6] == document_token[:6]
            ):
                matched.add(query_token)
                break
    return matched


def detected_query_domains(text: str) -> set[str]:
    query_terms = token_set(text)
    normalized = normalize_search_text(text)
    detected: set[str] = set()
    for row in (load_semantic_enrichment_contract().get("domains") or {}).values():
        domain = str(row.get("id") or "")
        for alias in row.get("aliases") or []:
            alias_normalized = normalize_search_text(str(alias))
            if alias_normalized and (alias_normalized in normalized or token_set(alias_normalized) & query_terms):
                detected.add(domain)
                break
    return detected


def field_overlap_score(query_terms: set[str], value: str | list[str], weight: int) -> tuple[int, bool]:
    text = " ".join(str(item) for item in value) if isinstance(value, list) else str(value or "")
    overlap = query_terms & token_set(text)
    return len(overlap) * weight, bool(overlap)


def query_cmd(args: argparse.Namespace) -> int:
    package_dir = Path(args.package_dir).resolve()
    query_tokens = token_set(args.text)
    query_concepts = concept_tokens(args.text)
    query_domains = detected_query_domains(args.text)
    fragments_path = package_dir / "fragments.jsonl"
    fragments_by_topic: dict[str, list[dict[str, Any]]] = defaultdict(list)
    if fragments_path.is_file():
        with fragments_path.open(encoding="utf-8") as fh:
            for line in fh:
                if line.strip():
                    fragment = json.loads(line)
                    fragments_by_topic[str(fragment.get("topic_uid"))].append(fragment)
    rows: list[tuple[int, dict[str, Any], dict[str, Any] | None, list[str]]] = []
    for line in (package_dir / "topics.jsonl").read_text(encoding="utf-8").split("\n"):
        if not line.strip():
            continue
        topic = json.loads(line)
        if args.audience and args.audience not in (topic.get("audiences") or []):
            continue
        if args.job and args.job not in (topic.get("jobs") or []):
            continue
        if topic.get("node_kind") not in {None, "content", "gap"} or topic.get("queryable") is False:
            continue
        if query_domains and not query_domains.intersection(set(topic.get("domains") or [])):
            continue
        content = (package_dir / topic["content_ref"]).read_text(encoding="utf-8")
        score = 0
        matched_fields: list[str] = []
        for field, weight in (("title", 14), ("aliases", 12), ("answers_questions", 11), ("summary", 6)):
            part, matched = field_overlap_score(query_tokens, topic.get(field) or "", weight)
            score += part
            if matched:
                matched_fields.append(field)
        best_fragment: dict[str, Any] | None = None
        best_fragment_score = 0
        fragments = fragments_by_topic.get(str(topic["uid"])) or markdown_fragments(
            str(topic["uid"]), str(topic.get("parent_uid") or ""), content, str(topic["content_ref"])
        )
        for fragment in fragments:
            fragment_score, fragment_matched = field_overlap_score(
                query_tokens,
                " ".join([str(fragment.get("heading") or ""), str(fragment.get("plain_text") or fragment.get("text") or "")]),
                4,
            )
            if fragment_matched and fragment_score > best_fragment_score:
                best_fragment = fragment
                best_fragment_score = fragment_score
        score += best_fragment_score
        searchable_text = " ".join(
            [
                str(topic.get("title") or ""),
                " ".join(str(value) for value in topic.get("aliases") or []),
                " ".join(str(value) for value in topic.get("answers_questions") or []),
                str(topic.get("summary") or ""),
                str((best_fragment or {}).get("plain_text") or (best_fragment or {}).get("text") or ""),
                content,
            ]
        )
        matched_concepts = matched_query_concepts(query_concepts, searchable_text)
        minimum_concepts = 1 if len(query_concepts) <= 1 else 2
        if len(matched_concepts) < minimum_concepts:
            continue
        matched_fields.append(f"concepts:{len(matched_concepts)}/{len(query_concepts)}")
        if query_domains:
            score += 20
            matched_fields.append("domain")
        if score >= 8:
            rows.append((score, topic, best_fragment, matched_fields))
    rows.sort(key=lambda item: (-item[0], item[1]["uid"]))
    if not rows:
        answer = {
            "answer_status": "gap",
            "release": (read_yaml(package_dir / "package.yaml") or {}).get("release", "unknown"),
            "answer": "В опубликованном пакете не найден подтвержденный ответ.",
            "applicability": [],
            "citations": [],
            "normative_levels": [],
            "next_step": "Создать editorial gap candidate; не формулировать новую норму автоматически.",
        }
    else:
        selected = rows[: min(args.limit, len(rows))]
        top_score, top_topic, top_fragment, _ = selected[0]
        top_is_gap = top_topic.get("node_kind") == "gap" or top_topic.get("coverage_status") == "gap"
        answer_text = ""
        if top_fragment:
            answer_text = str(top_fragment.get("text") or top_fragment.get("plain_text") or "")
        elif not top_is_gap:
            answer_text = (package_dir / top_topic["content_ref"]).read_text(encoding="utf-8")
        answer = {
            "answer_status": "gap" if top_is_gap else "documented",
            "release": (read_yaml(package_dir / "package.yaml") or {}).get("release", "unknown"),
            "answer": (
                "Раздел присутствует в стандарте, но не наполнен подтвержденным содержанием."
                if top_is_gap
                else answer_text[: args.max_chars].strip()
            ),
            "applicability": [x for x in [args.audience, args.job] if x],
            "detected_domains": sorted(query_domains),
            "citations": [
                {
                    "uid": topic["uid"],
                    "legacy_uids": topic.get("legacy_uids") or ([topic.get("storage_uid")] if topic.get("storage_uid") else []),
                    "title": topic["title"],
                    "fragment_uid": fragment.get("uid") if fragment else None,
                    "content_ref": topic["content_ref"],
                    "coverage_status": topic.get("coverage_status"),
                    "score": score,
                    "matched_fields": matched,
                }
                for score, topic, fragment, matched in selected
            ],
            "normative_levels": [],
            "next_step": "Создать editorial gap candidate; не формулировать новую норму автоматически." if top_is_gap else None,
        }
    print(json.dumps(answer, ensure_ascii=False, indent=2))
    return 0


def audit_cmd(args: argparse.Namespace) -> int:
    package_dir = Path(args.package_dir).resolve()
    project = read_yaml(Path(args.input).resolve()) or {}
    rules = [json.loads(line) for line in (package_dir / "rules.jsonl").read_text(encoding="utf-8").split("\n") if line.strip()]
    missing = [field for field in ("requirements", "components") if field not in project]
    result = {
        "audit_status": "needs_input" if missing or not rules else "pass",
        "release": (read_yaml(package_dir / "package.yaml") or {}).get("release", "unknown"),
        "checked_inputs": project,
        "applicable_rule_uids": [row["uid"] for row in rules],
        "violations": [],
        "missing_inputs": missing + (["approved_typed_rules"] if not rules else []),
        "citations": [{"uid": row["uid"], "statement": row.get("statement")} for row in rules],
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0 if result["audit_status"] == "pass" else 3


def markdown_table_count(content: str) -> int:
    count = 0
    for line in logical_markdown_lines(content):
        if not line.strip().startswith("|"):
            continue
        cells = markdown_table_cells(line)
        if markdown_is_separator(cells):
            count += 1
    return count


def migration_audit_cmd(args: argparse.Namespace) -> int:
    """Reconcile immutable Google baseline, migrated source, builds and agent package.

    This is a deterministic technical audit. It deliberately does not certify
    D8, which remains an independent human/agent acceptance gate.
    """
    baseline = Path(args.baseline).resolve()
    source_dir = Path(args.source_dir).resolve()
    html_path = Path(args.html).resolve() if args.html else None
    docx_path = Path(args.docx).resolve() if args.docx else None
    package_dir = Path(args.package_dir).resolve() if args.package_dir else None
    manifest = read_json(baseline / "baseline-manifest.json")
    document = read_json(baseline / "document.json")
    book, sections, topics, _ = collect_units(source_dir)
    section_by_tab = {str(row.get("source_tab_id")): row for row in sections}
    topic_by_uid = {str(row["uid"]): row for row in topics}
    asset_manifest = read_yaml(source_dir / "assets" / "manifest.yaml") or {}
    baseline_assets = {str(row.get("object_id")): row for row in manifest.get("assets") or []}
    source_assets = {str(row.get("object_id")): row for row in asset_manifest.get("assets") or []}
    asset_paths = {
        uid: f"../../assets/{Path(str(row.get('path') or '')).name}"
        for uid, row in source_assets.items()
        if row.get("status") == "downloaded"
    }
    checks: list[dict[str, Any]] = []

    def check(name: str, passed: bool, **details: Any) -> None:
        checks.append({"name": name, "status": "pass" if passed else "fail", **details})

    check(
        "baseline_identity",
        book.get("baseline_uid") == manifest.get("baseline_uid") and book.get("source_revision_id") == manifest.get("revision_id"),
        baseline_uid=manifest.get("baseline_uid"),
        revision_id=manifest.get("revision_id"),
    )
    tabs_checked = 0
    blocks_checked = 0
    topics_checked = 0
    baseline_tables = 0
    source_tables = 0
    tab_failures: list[str] = []
    block_failures: list[str] = []
    legacy_failures: list[str] = []
    topic_failures: list[str] = []
    for tab in iter_tabs(document.get("tabs") or []):
        props = tab_properties(tab)
        tab_id = str(props.get("tabId") or "unknown")
        title = str(props.get("title") or tab_id)
        is_buffer = title.strip().casefold() == "буфер"
        section = section_by_tab.get(tab_id)
        target = source_dir / "staging" / "buffer" if is_buffer else source_dir / "sections" / str((section or {}).get("uid") or "missing")
        source_tab_path = target / "source-tab.json"
        source_tab = read_json(source_tab_path) if source_tab_path.is_file() else None
        if source_tab != document_tab(tab):
            tab_failures.append(tab_id)
        tabs_checked += 1
        expected_blocks = body_blocks((document_tab(tab).get("body") or {}))
        actual_blocks = []
        blocks_path = target / "blocks.jsonl"
        if blocks_path.is_file():
            actual_blocks = [json.loads(line) for line in blocks_path.read_text(encoding="utf-8").split("\n") if line.strip()]
        if actual_blocks != expected_blocks:
            block_failures.append(tab_id)
        blocks_checked += len(expected_blocks)
        expected_markdown = "\n\n".join(x for x in (markdown_for_block(block, asset_paths) for block in expected_blocks) if x).strip() + "\n"
        legacy_path = target / "legacy.md"
        actual_markdown = legacy_path.read_text(encoding="utf-8") if legacy_path.is_file() else ""
        if actual_markdown != expected_markdown:
            legacy_failures.append(tab_id)
        baseline_tables += sum(1 for block in expected_blocks if block.get("type") == "table")
        source_tables += markdown_table_count(actual_markdown)
        if is_buffer or not section:
            continue
        topic_asset_paths = {key: value.replace("../../assets/", "../../../../assets/") for key, value in asset_paths.items()}
        expected_topics = split_topics(str(section["uid"]), title, expected_blocks, topic_asset_paths)
        if [row["uid"] for row in expected_topics] != list(section.get("topic_refs") or []):
            topic_failures.append(str(section["uid"]) + ":refs")
        for expected_topic in expected_topics:
            actual_topic = topic_by_uid.get(str(expected_topic["uid"]))
            if not actual_topic:
                topic_failures.append(str(expected_topic["uid"]) + ":missing")
                continue
            content_path = source_dir / str(actual_topic.get("content_ref") or "")
            actual_content = content_path.read_text(encoding="utf-8") if content_path.is_file() else ""
            if actual_content != expected_topic["content"] or actual_topic.get("title") != expected_topic["title"]:
                topic_failures.append(str(expected_topic["uid"]) + ":content")
            topics_checked += 1
    check("source_tabs_equal_google_baseline", not tab_failures, tabs=tabs_checked, failures=tab_failures)
    check("block_streams_rederived_exactly", not block_failures, blocks=blocks_checked, failures=block_failures)
    check("legacy_markdown_rederived_exactly", not legacy_failures, failures=legacy_failures)
    check("topic_split_and_content_rederived_exactly", not topic_failures, topics=topics_checked, failures=topic_failures)
    check(
        "table_inventory_preserved",
        baseline_tables == source_tables,
        baseline_tables=baseline_tables,
        source_tables=source_tables,
        note="Google soft line breaks are normalized only while parsing, not rewritten in the lossless source.",
    )

    asset_failures: list[str] = []
    for uid, baseline_row in baseline_assets.items():
        source_row = source_assets.get(uid)
        if not source_row or source_row.get("sha256") != baseline_row.get("sha256"):
            asset_failures.append(uid + ":manifest")
            continue
        path = source_dir / str(source_row.get("path") or "")
        if source_row.get("status") == "downloaded" and (not path.is_file() or sha256_file(path) != source_row.get("sha256")):
            asset_failures.append(uid + ":file")
    check(
        "asset_inventory_and_digests",
        not asset_failures and len(source_assets) == len(baseline_assets),
        baseline_assets=len(baseline_assets),
        source_assets=len(source_assets),
        inline=sum(1 for row in source_assets.values() if row.get("placement") == "inline"),
        positioned=sum(1 for row in source_assets.values() if row.get("placement") == "positioned"),
        failures=asset_failures,
    )

    semantic_uids = [str(row.get("semantic_uid") or "") for row in topics]
    node_counts: dict[str, int] = defaultdict(int)
    semantic_failures: list[str] = []
    for row in topics:
        semantic_uid = str(row.get("semantic_uid") or "")
        node_kind = str(row.get("node_kind") or "")
        node_counts[node_kind] += 1
        if not re.fullmatch(r"std_topic_[a-z0-9_]+", semantic_uid):
            semantic_failures.append(str(row.get("uid")) + ":semantic_uid")
        if not row.get("legacy_uids") or row.get("coverage_status") not in {"documented", "gap", "out_of_scope"}:
            semantic_failures.append(str(row.get("uid")) + ":metadata")
    check(
        "semantic_addressing",
        not semantic_failures and len(set(semantic_uids)) == len(semantic_uids),
        topics=len(topics),
        unique_semantic_uids=len(set(semantic_uids)),
        node_kinds=dict(sorted(node_counts.items())),
        failures=semantic_failures,
    )

    if html_path:
        html_text = html_path.read_text(encoding="utf-8") if html_path.is_file() else ""
        expected_published = [row for row in topics if row.get("node_kind") not in {"artifact", "attachment"} and row.get("publish") is not False]
        topic_anchor_count = len(re.findall(r"data-topic-uid=", html_text))
        html_table_count = len(re.findall(r"<table\b", html_text))
        html_image_count = len(re.findall(r"<img\b", html_text))
        missing_html_assets = [
            src for src in re.findall(r"<img[^>]+src=['\"]([^'\"]+)", html_text)
            if not (html_path.parent / src).is_file()
        ]
        html_ids = set(re.findall(r"\bid=['\"]([^'\"]+)", html_text))
        internal_link_targets = re.findall(r"\bhref=['\"]#([^'\"]+)", html_text)
        missing_internal_targets = sorted(set(internal_link_targets) - html_ids)
        html_toc_count = len(re.findall(r"\bid=['\"]std_toc['\"]", html_text))
        topic_navigation_count = len(re.findall(r"class=['\"][^'\"]*\btopic-navigation\b", html_text))
        main_match = re.search(r"<main\b[^>]*>(.*?)</main>", html_text, flags=re.S | re.I)
        main_html = main_match.group(1) if main_match else ""
        main_without_navigation = re.sub(r"<nav\b[^>]*>.*?</nav>", " ", main_html, flags=re.S | re.I)
        block_separated_html = re.sub(
            r"</(?:p|li|td|th|h[1-6]|tr|table|figure|figcaption|aside)>",
            " ",
            main_without_navigation,
            flags=re.I,
        )
        rendered_plain_text = parity_plain_text(re.sub(r"<[^>]+>", "", block_separated_html))
        checked_content_lines = 0
        missing_content_lines: list[dict[str, str]] = []
        for topic in topics:
            if topic.get("node_kind") == "artifact" or topic.get("publish") is False:
                continue
            content_path = source_dir / str(topic.get("content_ref") or "")
            if not content_path.is_file():
                continue
            for line in logical_markdown_lines(content_path.read_text(encoding="utf-8")):
                plain_line = parity_plain_text(markdown_plain_text(line).lstrip("- ").strip())
                if not plain_line or re.fullmatch(r"[-\s]+", plain_line):
                    continue
                checked_content_lines += 1
                if plain_line not in rendered_plain_text:
                    missing_content_lines.append({"topic_uid": topic_public_uid(topic), "text": plain_line[:240]})
        check(
            "normalized_html_structure",
            bool(html_text)
            and topic_anchor_count == len(expected_published)
            and html_table_count == baseline_tables
            and html_toc_count == 1
            and topic_navigation_count == len(expected_published)
            and not missing_internal_targets
            and not missing_content_lines
            and "#### " not in html_text
            and not contains_private_use(html_text)
            and not missing_html_assets,
            topic_anchors=topic_anchor_count,
            expected_topic_anchors=len(expected_published),
            tables=html_table_count,
            images=html_image_count,
            toc=html_toc_count,
            topic_navigation=topic_navigation_count,
            internal_links=len(internal_link_targets),
            missing_internal_targets=missing_internal_targets,
            checked_content_lines=checked_content_lines,
            missing_content_lines=missing_content_lines[:100],
            missing_assets=missing_html_assets,
        )
        expected_formatting = {"bold": 0, "italic": 0, "underline": 0, "external_links": 0}
        for section in sections:
            blocks_path = source_dir / "sections" / str(section["uid"]) / "blocks.jsonl"
            if not blocks_path.is_file():
                continue
            for raw_line in blocks_path.read_text(encoding="utf-8").split("\n"):
                if not raw_line.strip():
                    continue
                block = json.loads(raw_line)
                if (
                    block.get("type") != "paragraph"
                    or str(block.get("style") or "").startswith("HEADING_")
                ):
                    continue
                for run in block.get("runs") or []:
                    if not str(run.get("text") or "").strip():
                        continue
                    style = run.get("style") or {}
                    expected_formatting["bold"] += int(bool(style.get("bold")))
                    expected_formatting["italic"] += int(bool(style.get("italic")))
                    expected_formatting["underline"] += int(bool(style.get("underline")))
                    expected_formatting["external_links"] += int(bool((style.get("link") or {}).get("url")))
        rendered_formatting = {
            "bold": len(re.findall(r"<strong\b", main_html)),
            "italic": len(re.findall(r"<em\b", main_html)),
            "underline": len(re.findall(r"<u\b", main_html)),
            "external_links": len(re.findall(r"<a\b[^>]*href=['\"](?!#)", main_html)),
        }
        check(
            "normalized_html_inline_formatting",
            all(rendered_formatting[key] >= value for key, value in expected_formatting.items()),
            expected_source_segments=expected_formatting,
            rendered_tags=rendered_formatting,
            note="Headings are governed by structural rules; generated warnings may add formatting tags.",
        )

    if docx_path:
        docx_tables = 0
        docx_images = 0
        docx_bookmarks = 0
        docx_raw_hash_markers = 0
        docx_private_use = 0
        docx_internal_links = 0
        docx_missing_internal_targets: list[str] = []
        docx_error = None
        try:
            with zipfile.ZipFile(docx_path) as archive:
                document_xml = ET.fromstring(archive.read("word/document.xml"))
                word_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
                docx_tables = sum(1 for _ in document_xml.iter(word_ns + "tbl"))
                docx_images = sum(1 for _ in document_xml.iter(word_ns + "drawing"))
                docx_bookmarks = sum(1 for _ in document_xml.iter(word_ns + "bookmarkStart"))
                bookmark_names = {
                    str(node.get(word_ns + "name") or "")
                    for node in document_xml.iter(word_ns + "bookmarkStart")
                }
                hyperlink_targets = [
                    str(node.get(word_ns + "anchor") or "")
                    for node in document_xml.iter(word_ns + "hyperlink")
                    if node.get(word_ns + "anchor")
                ]
                docx_internal_links = len(hyperlink_targets)
                docx_missing_internal_targets = sorted(set(hyperlink_targets) - bookmark_names)
                text_nodes = [str(node.text or "") for node in document_xml.iter(word_ns + "t")]
                docx_raw_hash_markers = sum("#### " in value for value in text_nodes)
                docx_private_use = sum(contains_private_use(value) for value in text_nodes)
        except Exception as exc:  # noqa: BLE001
            docx_error = f"{type(exc).__name__}: {exc}"
        published_tab_ids = {str(section.get("source_tab_id")) for section in sections}
        expected_docx_images = sum(
            1
            for row in source_assets.values()
            if row.get("status") == "downloaded" and str(row.get("tab_id")) in published_tab_ids
        )
        check(
            "normalized_docx_structure",
            docx_error is None
            and docx_tables == baseline_tables
            and docx_images == expected_docx_images
            and docx_bookmarks >= len(sections) + len(
                [row for row in topics if row.get("node_kind") not in {"artifact", "attachment"} and row.get("publish") is not False]
            )
            and docx_internal_links > 0
            and not docx_missing_internal_targets
            and docx_raw_hash_markers == 0
            and docx_private_use == 0,
            tables=docx_tables,
            images=docx_images,
            expected_images=expected_docx_images,
            bookmarks=docx_bookmarks,
            internal_links=docx_internal_links,
            missing_internal_targets=docx_missing_internal_targets,
            raw_hash_markers=docx_raw_hash_markers,
            private_use_text_nodes=docx_private_use,
            error=docx_error,
        )

    if package_dir:
        package = read_yaml(package_dir / "package.yaml") if (package_dir / "package.yaml").is_file() else {}
        package_topics = [json.loads(line) for line in (package_dir / "topics.jsonl").read_text(encoding="utf-8").split("\n") if line.strip()] if (package_dir / "topics.jsonl").is_file() else []
        package_uids = [str(row.get("uid")) for row in package_topics]
        package_failures: list[str] = []
        for row in package_topics:
            content_path = package_dir / str(row.get("content_ref") or "")
            if not content_path.is_file():
                package_failures.append(str(row.get("uid")) + ":content")
                continue
            for ref in re.findall(r"!\[[^]]*\]\(([^)]+)\)", content_path.read_text(encoding="utf-8")):
                if not (content_path.parent / ref).resolve().is_file():
                    package_failures.append(str(row.get("uid")) + ":asset")
        expected_package_topics = sum(1 for row in topics if row.get("node_kind") not in {"artifact", "attachment"} and row.get("publish") is not False)
        check(
            "agent_package_integrity",
            bool(package)
            and len(package_topics) == expected_package_topics
            and len(package_uids) == len(set(package_uids))
            and package.get("buffer_included") is False
            and not package_failures,
            topics=len(package_topics),
            expected_topics=expected_package_topics,
            fragments=(package.get("counts") or {}).get("fragments"),
            assets=(package.get("counts") or {}).get("assets"),
            buffer_included=package.get("buffer_included"),
            failures=package_failures,
        )

    passed = all(row["status"] == "pass" for row in checks)
    severity_by_check = {
        "baseline_identity": "critical",
        "source_tabs_equal_google_baseline": "critical",
        "block_streams_rederived_exactly": "critical",
        "legacy_markdown_rederived_exactly": "critical",
        "topic_split_and_content_rederived_exactly": "critical",
        "table_inventory_preserved": "critical",
        "asset_inventory_and_digests": "critical",
        "semantic_addressing": "critical",
        "normalized_html_structure": "critical",
        "normalized_html_inline_formatting": "material",
        "normalized_docx_structure": "critical",
        "agent_package_integrity": "critical",
    }
    findings = [
        {
            "severity": severity_by_check.get(row["name"], "material"),
            "check_id": row["name"],
            "details": {key: value for key, value in row.items() if key not in {"name", "status"}},
        }
        for row in checks
        if row["status"] != "pass"
    ]
    report = {
        "schema_version": "1.0",
        "audit_type": "deterministic_migration_and_agent_access",
        "generated_at": utc_now(),
        "status": "pass_with_known_manual_gate" if passed else "fail",
        "d8_independent_certification": "deferred_not_self_certified",
        "baseline_uid": manifest.get("baseline_uid"),
        "source_revision_id": manifest.get("revision_id"),
        "severity_model": ["critical", "material", "cosmetic"],
        "findings_summary": {
            severity: sum(1 for finding in findings if finding["severity"] == severity)
            for severity in ("critical", "material", "cosmetic")
        },
        "findings": findings,
        "checks": checks,
    }
    if args.output:
        write_json(Path(args.output).resolve(), report)
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0 if passed else 2


def inventory_cmd(args: argparse.Namespace) -> int:
    baseline = Path(args.baseline).resolve()
    manifest = read_json(baseline / "baseline-manifest.json")
    doc = read_json(baseline / "document.json")
    totals = {"tabs": 0, "paragraphs": 0, "tables": 0, "characters": 0, "inline_images": 0, "positioned_images": 0}
    tabs_out: list[dict[str, Any]] = []
    for tab in iter_tabs(doc.get("tabs") or []):
        props = tab_properties(tab)
        blocks = body_blocks((document_tab(tab).get("body") or {}))
        row = {
            "tab_id": props.get("tabId"),
            "title": props.get("title"),
            "paragraphs": sum(1 for x in blocks if x["type"] == "paragraph"),
            "tables": sum(1 for x in blocks if x["type"] == "table"),
            "characters": sum(len(x.get("text") or "") for x in blocks if x["type"] == "paragraph"),
            "inline_images": len(document_tab(tab).get("inlineObjects") or {}),
            "positioned_images": len(document_tab(tab).get("positionedObjects") or {}),
        }
        tabs_out.append(row)
        totals["tabs"] += 1
        for key in totals:
            if key != "tabs":
                totals[key] += int(row.get(key) or 0)
    report = {"baseline_uid": manifest["baseline_uid"], "revision_id": manifest["revision_id"], "totals": totals, "tabs": tabs_out}
    if args.output:
        write_json(Path(args.output).resolve(), report)
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0


def diff_cmd(args: argparse.Namespace) -> int:
    before = read_json(Path(args.from_manifest).resolve())
    after = read_json(Path(args.to_manifest).resolve())
    diff = {
        "from": before.get("baseline_uid") or before.get("package_uid"),
        "to": after.get("baseline_uid") or after.get("package_uid"),
        "changed": before != after,
        "from_digest": sha256_bytes(json.dumps(before, sort_keys=True).encode()),
        "to_digest": sha256_bytes(json.dumps(after, sort_keys=True).encode()),
    }
    print(json.dumps(diff, ensure_ascii=False, indent=2))
    return 0


def release_candidate_cmd(args: argparse.Namespace) -> int:
    if not re.fullmatch(r"\d+\.\d+\.\d+", args.version):
        raise SystemExit("--version must be SemVer X.Y.Z")
    source_dir = Path(args.source_dir).resolve()
    output_dir = Path(args.output_dir).resolve()
    if output_dir.exists() and any(output_dir.iterdir()) and not args.force:
        raise SystemExit(f"Candidate directory is not empty: {output_dir}. Use --force for deliberate regeneration.")
    output_dir.mkdir(parents=True, exist_ok=True)
    book, sections, topics, rules = collect_units(source_dir)
    units = [book, *sections, *topics, *rules]
    manifest_rows = [
        {
            "uid": unit["uid"],
            "type": unit["type"],
            "revision": unit.get("revision"),
            "digest": unit.get("digest"),
            "status": unit.get("status"),
        }
        for unit in units
    ]
    manifest_bytes = json.dumps(manifest_rows, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode("utf-8")
    included_changesets = list(dict.fromkeys(args.include_changeset or []))
    release = {
        "uid": stable_uid("std_release", args.version),
        "type": "release",
        "status": "draft",
        "revision": 1,
        "digest": None,
        "privacy": "internal",
        "source_refs": list(book.get("source_refs") or []),
        "change_refs": included_changesets,
        "introduced_in": None,
        "last_changed_in": None,
        "allowed_outputs": ["internal_review"],
        "version": args.version,
        "included_changesets": included_changesets,
        "manifest_digest": sha256_bytes(manifest_bytes),
        "published_at": None,
    }
    write_json(output_dir / "content-manifest.json", manifest_rows)
    write_yaml(output_dir / "release.yaml", release)
    atomic_write_text(
        output_dir / "CHANGELOG.md",
        f"# Release candidate {args.version}\n\n"
        "Status: internal review; not published.\n\n"
        f"- Sections: {len(sections)}\n- Topics: {len(topics)}\n- Typed rules: {len(rules)}\n"
        f"- Included changesets: {len(included_changesets)}\n- Manifest: `{release['manifest_digest']}`\n",
    )
    result = {"candidate": str(output_dir), "release": release, "counts": {"sections": len(sections), "topics": len(topics), "rules": len(rules)}}
    print(json.dumps(result, ensure_ascii=False))
    return 0


def parser() -> argparse.ArgumentParser:
    ap = argparse.ArgumentParser(prog="standard_book", description=__doc__)
    sub = ap.add_subparsers(dest="command", required=True)

    p = sub.add_parser("capture-google", help="Read Google Docs API and create immutable baseline")
    p.add_argument("--document-id", required=True)
    p.add_argument("--output", required=True)
    p.add_argument("--google-mcp-root")
    p.add_argument("--force", action="store_true")
    p.add_argument("--workers", type=int, default=8, help="Parallel image downloads (default: 8)")
    p.set_defaults(func=capture_google)

    p = sub.add_parser("inventory")
    p.add_argument("--baseline", required=True)
    p.add_argument("--output")
    p.set_defaults(func=inventory_cmd)

    p = sub.add_parser("extract")
    p.add_argument("--baseline", required=True)
    p.add_argument("--source-dir", default=str(DEFAULT_SOURCE_DIR))
    p.add_argument("--manifest-candidate", default=str(DEFAULT_MANIFEST_CANDIDATE))
    p.add_argument("--document-id")
    p.add_argument("--force", action="store_true", help="Deliberately regenerate the imported baseline tree")
    p.set_defaults(func=extract)

    p = sub.add_parser("validate")
    p.add_argument("--source-dir", default=str(DEFAULT_SOURCE_DIR))
    p.add_argument("--output")
    p.set_defaults(func=validate)

    p = sub.add_parser("remediate-migration", help="Add stable semantic addressing and classify migrated topic nodes")
    p.add_argument("--source-dir", default=str(DEFAULT_SOURCE_DIR))
    p.add_argument("--contract", default=str(SEMANTIC_ENRICHMENT_CONTRACT))
    p.add_argument("--output")
    p.set_defaults(func=remediate_migration_cmd)

    p = sub.add_parser("build")
    p.add_argument("--source-dir", default=str(DEFAULT_SOURCE_DIR))
    p.add_argument("--output-dir", default=str(DEFAULT_BUILD_DIR))
    p.add_argument("--profile", choices=["legacy-fidelity", "standard-normalized"], required=True)
    p.add_argument("--format", choices=["html", "docx", "all"], default="all")
    p.add_argument("--section", action="append", dest="sections", help="Build only this section UID; repeat for multiple sections")
    p.set_defaults(func=build_cmd)

    p = sub.add_parser("index")
    p.add_argument("--source-dir", default=str(DEFAULT_SOURCE_DIR))
    p.add_argument("--output-dir", required=True)
    p.add_argument("--release", default="0.0.0-baseline")
    p.add_argument("--force", action="store_true")
    p.set_defaults(func=index_cmd)

    p = sub.add_parser("query")
    p.add_argument("--package-dir", required=True)
    p.add_argument("--text", required=True)
    p.add_argument("--audience")
    p.add_argument("--job")
    p.add_argument("--limit", type=int, default=3)
    p.add_argument("--max-chars", type=int, default=1600)
    p.set_defaults(func=query_cmd)

    p = sub.add_parser("audit")
    p.add_argument("--package-dir", required=True)
    p.add_argument("--input", required=True)
    p.set_defaults(func=audit_cmd)

    p = sub.add_parser("audit-migration", help="Reconcile baseline, migrated source, normalized build and agent package; never self-certifies D8")
    p.add_argument("--baseline", required=True)
    p.add_argument("--source-dir", default=str(DEFAULT_SOURCE_DIR))
    p.add_argument("--html")
    p.add_argument("--docx")
    p.add_argument("--package-dir")
    p.add_argument("--output")
    p.set_defaults(func=migration_audit_cmd)

    p = sub.add_parser("diff")
    p.add_argument("--from-manifest", required=True)
    p.add_argument("--to-manifest", required=True)
    p.set_defaults(func=diff_cmd)

    p = sub.add_parser("release-candidate", help="Create an internal immutable release candidate manifest; never publishes")
    p.add_argument("--source-dir", default=str(DEFAULT_SOURCE_DIR))
    p.add_argument("--output-dir", required=True)
    p.add_argument("--version", required=True)
    p.add_argument("--include-changeset", action="append", default=[])
    p.add_argument("--force", action="store_true")
    p.set_defaults(func=release_candidate_cmd)

    p = sub.add_parser("propose-rules", help="Extract lexical rule candidates for expert review; never approves or publishes them")
    p.add_argument("--source-dir", default=str(DEFAULT_SOURCE_DIR))
    p.add_argument("--section", required=True)
    p.add_argument("--output", required=True)
    p.set_defaults(func=propose_rules_cmd)
    return ap


def main() -> int:
    if sys.platform == "win32":
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    args = parser().parse_args()
    return int(args.func(args))


if __name__ == "__main__":
    raise SystemExit(main())
